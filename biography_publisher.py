import streamlit as st
from datetime import datetime
import io
import base64
import os
import re
import html
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from fpdf import FPDF
import ebooklib
from ebooklib import epub
import tempfile

def clean_text(text):
    """Convert HTML entities to regular characters"""
    if not text:
        return text
    
    # Convert &nbsp; to space
    text = text.replace('&nbsp;', ' ')
    
    # Handle other common HTML entities
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&quot;', '"')
    text = text.replace('&#39;', "'")
    
    # Remove HTML tags but keep paragraph structure
    text = re.sub(r'<[^>]+>', '', text)
    
    return text.strip()

def show_celebration():
    """Show a celebration animation when book is generated"""
    st.balloons()
    st.success("🎉 Your book has been generated successfully!")

def generate_docx(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate a Word document from stories"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set document styling
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.first_line_indent = Inches(0.25)
    
    # COVER PAGE
    if cover_choice == "uploaded" and cover_image:
        try:
            image_stream = io.BytesIO(cover_image)
            doc.add_picture(image_stream, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add title and author below image
            cover_para = doc.add_paragraph()
            cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_run = cover_para.add_run(title)
            cover_run.font.size = Pt(42)
            cover_run.font.bold = True
            
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"by {author}")
            author_run.font.size = Pt(24)
            author_run.font.italic = True
            
            doc.add_page_break()
        except Exception as e:
            # Fallback to simple text cover
            cover_para = doc.add_paragraph()
            cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_run = cover_para.add_run(title)
            cover_run.font.size = Pt(42)
            cover_run.font.bold = True
            
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"by {author}")
            author_run.font.size = Pt(24)
            author_run.font.italic = True
            
            doc.add_page_break()
    else:
        # Simple text cover
        cover_para = doc.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover_para.add_run(title)
        cover_run.font.size = Pt(42)
        cover_run.font.bold = True
        
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"by {author}")
        author_run.font.size = Pt(24)
        author_run.font.italic = True
        
        doc.add_page_break()
    
    # Add publication info
    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_para.add_run(f"© {datetime.now().year} {author}. All rights reserved.")
    doc.add_paragraph()
    
    # Table of Contents
    if include_toc:
        doc.add_page_break()
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run("Table of Contents")
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_para.paragraph_format.space_after = Pt(12)
        
        # Group by session for TOC
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        for session_title in sessions.keys():
            p = doc.add_paragraph()
            p.add_run(f"  {session_title}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add stories
    doc.add_page_break()
    
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        
        # Add session header if new session
        if session_title != current_session:
            current_session = session_title
            session_para = doc.add_paragraph()
            session_run = session_para.add_run(session_title)
            session_run.font.size = Pt(16)
            session_run.font.bold = True
            session_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            session_para.paragraph_format.space_before = Pt(12)
            session_para.paragraph_format.space_after = Pt(6)
        
        if format_style == "interview":
            # Add question
            question_text = story.get('question', '')
            clean_question = clean_text(question_text)
            if clean_question:
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(clean_question)
                q_run.font.bold = True
                q_run.font.italic = True
                q_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                q_para.paragraph_format.space_before = Pt(6)
                q_para.paragraph_format.space_after = Pt(3)
        
        # Add answer
        answer_text = story.get('answer_text', '')
        if answer_text:
            clean_answer = clean_text(answer_text)
            
            # Split into paragraphs and add
            paragraphs = clean_answer.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(6)
        
        # Add images if any
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    try:
                        img_data = base64.b64decode(img['base64'])
                        img_stream = io.BytesIO(img_data)
                        
                        # Add image centered
                        doc.add_picture(img_stream, width=Inches(4))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add caption
                        if img.get('caption'):
                            caption_para = doc.add_paragraph()
                            clean_caption = clean_text(img['caption'])
                            caption_run = caption_para.add_run(clean_caption)
                            caption_run.font.size = Pt(10)
                            caption_run.font.italic = True
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.paragraph_format.space_before = Pt(3)
                            caption_para.paragraph_format.space_after = Pt(6)
                    except Exception as e:
                        pass
        
        # Add spacing between stories
        doc.add_paragraph()
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

def generate_html(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate an HTML document from stories"""
    
    html_parts = []
    
    # HTML header with styling
    html_parts.append(f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{html.escape(title)}</title>
    <style>
        body {{
            font-family: 'Georgia', serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            background: #fff;
        }}
        h1 {{
            font-size: 42px;
            text-align: center;
            margin-bottom: 10px;
            color: #000;
        }}
        h2 {{
            font-size: 28px;
            margin-top: 40px;
            margin-bottom: 20px;
            color: #444;
            border-bottom: 2px solid #eee;
            padding-bottom: 10px;
        }}
        .author {{
            text-align: center;
            font-size: 18px;
            color: #666;
            margin-bottom: 40px;
            font-style: italic;
        }}
        .question {{
            font-weight: bold;
            font-size: 18px;
            margin-top: 30px;
            margin-bottom: 10px;
            color: #2c3e50;
            border-left: 4px solid #3498db;
            padding-left: 15px;
        }}
        .story-image {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border-radius: 5px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .image-caption {{
            text-align: center;
            font-size: 14px;
            color: #666;
            margin-top: -10px;
            margin-bottom: 20px;
            font-style: italic;
        }}
        .cover-page {{
            text-align: center;
            margin-bottom: 50px;
            page-break-after: always;
            min-height: 80vh;
            display: flex;
            flex-direction: column;
            justify-content: center;
        }}
        .cover-image {{
            max-width: 100%;
            max-height: 70vh;
            object-fit: contain;
            margin: 20px auto;
            border-radius: 10px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }}
        .simple-cover {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 60px 20px;
            border-radius: 10px;
            color: white;
            margin: 20px;
        }}
        .simple-cover h1 {{
            color: white;
        }}
        .simple-cover .author {{
            color: rgba(255,255,255,0.9);
        }}
        .copyright {{
            text-align: center;
            font-size: 12px;
            color: #999;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #eee;
        }}
        .toc {{
            background: #f9f9f9;
            padding: 20px;
            border-radius: 5px;
            margin: 30px 0;
        }}
        .toc ul {{
            list-style-type: none;
            padding-left: 0;
        }}
        .toc li {{
            margin-bottom: 10px;
        }}
        .toc a {{
            color: #3498db;
            text-decoration: none;
        }}
        .toc a:hover {{
            text-decoration: underline;
        }}
        @media print {{
            body {{
                padding: 0.5in;
            }}
            .cover-page {{
                page-break-after: always;
                min-height: auto;
            }}
        }}
    </style>
</head>
<body>
""")
    
    # COVER PAGE
    html_parts.append('<div class="cover-page">')
    
    if cover_choice == "uploaded" and cover_image:
        try:
            img_base64 = base64.b64encode(cover_image).decode()
            html_parts.append(f'''
                <img src="data:image/jpeg;base64,{img_base64}" class="cover-image">
                <h1>{html.escape(title)}</h1>
                <p class="author">by {html.escape(author)}</p>
            ''')
        except Exception:
            html_parts.append(f'''
                <div class="simple-cover">
                    <h1>{html.escape(title)}</h1>
                    <p class="author">by {html.escape(author)}</p>
                </div>
            ''')
    else:
        html_parts.append(f'''
            <div class="simple-cover">
                <h1>{html.escape(title)}</h1>
                <p class="author">by {html.escape(author)}</p>
            </div>
        ''')
    
    html_parts.append('</div>')
    
    # Copyright page
    html_parts.append(f'<p class="copyright">© {datetime.now().year} {html.escape(author)}. All rights reserved.</p>')
    
    # Table of Contents
    if include_toc:
        html_parts.append('<div class="toc">')
        html_parts.append('<h3>Table of Contents</h3>')
        html_parts.append('<ul>')
        
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        for session_title in sessions.keys():
            anchor = re.sub(r'[^a-zA-Z0-9\s-]', '', session_title.lower().replace(' ', '-'))
            html_parts.append(f'<li><a href="#{anchor}">{html.escape(session_title)}</a></li>')
        
        html_parts.append('</ul>')
        html_parts.append('</div>')
    
    # Add stories
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        anchor = re.sub(r'[^a-zA-Z0-9\s-]', '', session_title.lower().replace(' ', '-'))
        
        if session_title != current_session:
            current_session = session_title
            html_parts.append(f'<h2 id="{anchor}">{html.escape(session_title)}</h2>')
        
        if format_style == "interview":
            question_text = story.get('question', '')
            if question_text:
                clean_question = clean_text(question_text)
                html_parts.append(f'<div class="question">{html.escape(clean_question)}</div>')
        
        # Format answer
        answer_text = story.get('answer_text', '')
        if answer_text:
            clean_answer = clean_text(answer_text)
            
            paragraphs = clean_answer.split('\n')
            for para in paragraphs:
                if para.strip():
                    html_parts.append(f'<p>{html.escape(para.strip())}</p>')
        
        # Add images
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    html_parts.append(f'<img src="data:image/jpeg;base64,{img["base64"]}" class="story-image">')
                    if img.get('caption'):
                        clean_caption = clean_text(img['caption'])
                        html_parts.append(f'<p class="image-caption">{html.escape(clean_caption)}</p>')
        
        html_parts.append('<hr style="margin: 30px 0; border: none; border-top: 1px dashed #ccc;">')
    
    html_parts.append("""
</body>
</html>""")
    
    return '\n'.join(html_parts)

def generate_epub(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate a Reflowable EPUB file"""
    
    book = epub.EpubBook()
    
    # Set metadata
    book.set_identifier('id123456')
    book.set_title(title)
    book.set_language('en')
    book.add_author(author)
    
    # Create cover
    if cover_choice == "uploaded" and cover_image:
        book.set_cover("cover.jpg", cover_image)
    
    # Create CSS style
    style = '''
    body {
        font-family: Georgia, serif;
        line-height: 1.6;
        margin: 2em;
    }
    h1 {
        font-size: 2em;
        text-align: center;
        margin-top: 2em;
    }
    h2 {
        font-size: 1.5em;
        color: #444;
        border-bottom: 1px solid #ddd;
        padding-bottom: 0.3em;
    }
    .question {
        font-weight: bold;
        font-style: italic;
        color: #2c3e50;
        margin-top: 1.5em;
    }
    .story-image {
        max-width: 100%;
        height: auto;
        display: block;
        margin: 1em auto;
    }
    .image-caption {
        text-align: center;
        font-size: 0.9em;
        color: #666;
        font-style: italic;
    }
    .cover-page {
        text-align: center;
        margin-top: 30%;
    }
    .copyright {
        text-align: center;
        font-size: 0.8em;
        color: #999;
        margin-top: 2em;
    }
    '''
    
    nav_css = epub.EpubItem(
        uid="style_nav",
        file_name="style/nav.css",
        media_type="text/css",
        content=style
    )
    book.add_item(nav_css)
    
    chapters = []
    spine = ['nav']
    
    # Create cover page
    cover_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
<head>
    <link rel="stylesheet" type="text/css" href="style/nav.css"/>
</head>
<body>
    <div class="cover-page">
        <h1>{html.escape(title)}</h1>
        <h3>by {html.escape(author)}</h3>
        <p class="copyright">© {datetime.now().year}</p>
    </div>
</body>
</html>
'''
    
    cover_page = epub.EpubHtml(
        title='Cover',
        file_name='cover.xhtml',
        lang='en'
    )
    cover_page.content = cover_content
    cover_page.add_item(nav_css)
    book.add_item(cover_page)
    spine.append(cover_page)
    
    # Create TOC
    if include_toc:
        toc_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
<head>
    <link rel="stylesheet" type="text/css" href="style/nav.css"/>
</head>
<body>
    <h1>Table of Contents</h1>
    <ul>
'''
        
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        for session_title in sessions.keys():
            anchor = re.sub(r'[^a-zA-Z0-9\s-]', '', session_title.lower().replace(' ', '-'))
            toc_content += f'<li><a href="{anchor}.xhtml">{html.escape(session_title)}</a></li>\n'
        
        toc_content += '''
    </ul>
</body>
</html>
'''
        
        toc_page = epub.EpubHtml(
            title='Table of Contents',
            file_name='toc.xhtml',
            lang='en'
        )
        toc_page.content = toc_content
        toc_page.add_item(nav_css)
        book.add_item(toc_page)
        spine.append(toc_page)
    
    # Create story chapters
    current_session = None
    chapter_index = 1
    chapter_list = []
    
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        
        if session_title != current_session:
            current_session = session_title
            chapter_file = re.sub(r'[^a-zA-Z0-9\s-]', '', session_title.lower().replace(' ', '-')) + '.xhtml'
            
            chapter_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
<head>
    <link rel="stylesheet" type="text/css" href="style/nav.css"/>
</head>
<body>
    <h1>{html.escape(session_title)}</h1>
'''
            
            # Create chapter
            chapter = epub.EpubHtml(
                title=session_title,
                file_name=chapter_file,
                lang='en'
            )
            chapter_list.append(chapter)
            chapter_index += 1
        
        # Add story content to current chapter
        if format_style == "interview":
            question_text = story.get('question', '')
            if question_text:
                clean_question = clean_text(question_text)
                chapter_content += f'<div class="question">{html.escape(clean_question)}</div>\n'
        
        answer_text = story.get('answer_text', '')
        if answer_text:
            clean_answer = clean_text(answer_text)
            paragraphs = clean_answer.split('\n')
            for para in paragraphs:
                if para.strip():
                    chapter_content += f'<p>{html.escape(para.strip())}</p>\n'
        
        # Add images
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    img_data = base64.b64decode(img['base64'])
                    img_file = f"image_{chapter_index}_{hash(img_data) % 10000}.jpg"
                    
                    # Add image to book
                    book_image = epub.EpubImage()
                    book_image.file_name = f"images/{img_file}"
                    book_image.media_type = "image/jpeg"
                    book_image.content = img_data
                    book.add_item(book_image)
                    
                    chapter_content += f'<img src="images/{img_file}" class="story-image"/>\n'
                    
                    if img.get('caption'):
                        clean_caption = clean_text(img['caption'])
                        chapter_content += f'<p class="image-caption">{html.escape(clean_caption)}</p>\n'
        
        chapter_content += '<hr/>\n</body>\n</html>'
        chapter.content = chapter_content
        chapter.add_item(nav_css)
        book.add_item(chapter)
        spine.append(chapter)
    
    # Add navigation
    book.toc = (
        [epub.Section('Cover'), cover_page] +
        ([epub.Section('Table of Contents'), toc_page] if include_toc else []) +
        [epub.Section(session_title, chapter) for session_title, chapter in zip(sessions.keys(), chapter_list)]
    )
    
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    book.spine = spine
    
    # Save to bytes
    epub_bytes = io.BytesIO()
    epub.write_epub(epub_bytes, book)
    epub_bytes.seek(0)
    
    return epub_bytes.getvalue()

class PDF(FPDF):
    def header(self):
        if self.page_no() > 1:  # Skip header on cover page
            self.set_font('Times', 'I', 8)
            self.cell(0, 10, '', 0, 0, 'R')
            self.ln(20)
    
    def footer(self):
        if self.page_no() > 1:  # Skip footer on cover page
            self.set_y(-15)
            self.set_font('Times', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate a PDF/X compatible PDF file using FPDF2"""
    
    pdf = PDF()
    pdf.add_page()
    
    # Set margins (1 inch = 25.4 mm)
    pdf.set_left_margin(25.4)
    pdf.set_right_margin(25.4)
    pdf.set_top_margin(25.4)
    pdf.set_auto_page_break(True, margin=25.4)
    
    # Set default font
    pdf.set_font('Times', '', 12)
    
    # COVER PAGE
    if cover_choice == "uploaded" and cover_image:
        try:
            # Save cover image temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                tmp_file.write(cover_image)
                tmp_path = tmp_file.name
            
            # Add image centered
            pdf.image(tmp_path, x=50, y=50, w=110)
            os.unlink(tmp_path)
            
            # Add title and author below image
            pdf.ln(150)
            pdf.set_font('Times', 'B', 42)
            pdf.cell(0, 20, title, 0, 1, 'C')
            pdf.set_font('Times', 'I', 24)
            pdf.cell(0, 20, f'by {author}', 0, 1, 'C')
        except:
            pdf.set_font('Times', 'B', 42)
            pdf.cell(0, 100, title, 0, 1, 'C')
            pdf.set_font('Times', 'I', 24)
            pdf.cell(0, 20, f'by {author}', 0, 1, 'C')
    else:
        pdf.set_font('Times', 'B', 42)
        pdf.cell(0, 100, title, 0, 1, 'C')
        pdf.set_font('Times', 'I', 24)
        pdf.cell(0, 20, f'by {author}', 0, 1, 'C')
    
    # Copyright page
    pdf.add_page()
    pdf.set_font('Times', '', 12)
    pdf.ln(100)
    pdf.cell(0, 10, f'© {datetime.now().year} {author}. All rights reserved.', 0, 1, 'C')
    
    # Table of Contents
    if include_toc:
        pdf.add_page()
        pdf.set_font('Times', 'B', 18)
        pdf.cell(0, 20, 'Table of Contents', 0, 1, 'C')
        pdf.ln(10)
        
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        pdf.set_font('Times', '', 12)
        for session_title in sessions.keys():
            pdf.cell(0, 10, f'  {session_title}', 0, 1)
    
    # Stories
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        
        if session_title != current_session:
            current_session = session_title
            pdf.add_page()
            pdf.set_font('Times', 'B', 16)
            pdf.cell(0, 20, session_title, 0, 1, 'C')
            pdf.ln(10)
        
        if format_style == "interview":
            question_text = story.get('question', '')
            if question_text:
                clean_question = clean_text(question_text)
                pdf.set_font('Times', 'BI', 12)
                pdf.multi_cell(0, 6, clean_question)
                pdf.ln(5)
        
        answer_text = story.get('answer_text', '')
        if answer_text:
            clean_answer = clean_text(answer_text)
            pdf.set_font('Times', '', 12)
            paragraphs = clean_answer.split('\n')
            for para in paragraphs:
                if para.strip():
                    pdf.multi_cell(0, 6, para.strip())
                    pdf.ln(3)
        
        # Add images
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    try:
                        img_data = base64.b64decode(img['base64'])
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                            tmp_file.write(img_data)
                            tmp_path = tmp_file.name
                        
                        # Calculate image width to fit page margins
                        page_width = pdf.w - 2 * pdf.l_margin
                        pdf.image(tmp_path, x=pdf.l_margin + 10, w=page_width - 20)
                        os.unlink(tmp_path)
                        
                        if img.get('caption'):
                            clean_caption = clean_text(img['caption'])
                            pdf.set_font('Times', 'I', 10)
                            pdf.ln(5)
                            pdf.cell(0, 5, clean_caption, 0, 1, 'C')
                    except:
                        pass
        
        pdf.ln(10)
    
    # Output as bytes
    return pdf.output(dest='S').encode('latin-1')

def generate_rtf(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate an RTF file"""
    
    rtf_parts = []
    
    # RTF Header
    rtf_parts.append("""{\\rtf1\\ansi\\deff0
{\\fonttbl{\\f0\\fnil\\fcharset0 Times New Roman;}}
\\viewkind4\\uc1\\pard\\sa200\\sl276\\slmult1\\f0\\fs24
""")
    
    # Cover page
    if cover_choice == "uploaded" and cover_image:
        rtf_parts.append("\\par\\pard\\qc ")
        rtf_parts.append(f"\\b\\fs48 {title}\\b0\\par\\par")
        rtf_parts.append(f"\\i\\fs36 by {author}\\i0\\par\\par")
        rtf_parts.append("\\par [Cover Image] \\par\\par")
    else:
        rtf_parts.append("\\par\\pard\\qc ")
        rtf_parts.append(f"\\b\\fs48 {title}\\b0\\par\\par")
        rtf_parts.append(f"\\i\\fs36 by {author}\\i0\\par\\par")
    
    rtf_parts.append("\\par\\pard\\qc ")
    rtf_parts.append(f"\\fs20 © {datetime.now().year} {author}. All rights reserved.\\par")
    rtf_parts.append("\\page\\par")
    
    # Table of Contents
    if include_toc:
        rtf_parts.append("\\pard\\qc \\b\\fs36 Table of Contents\\b0\\par\\par")
        
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        for session_title in sessions.keys():
            rtf_parts.append(f"\\pard\\li720 {session_title}\\par")
        
        rtf_parts.append("\\page\\par")
    
    # Stories
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        
        if session_title != current_session:
            current_session = session_title
            rtf_parts.append(f"\\pard\\qc \\b\\fs32 {session_title}\\b0\\par\\par")
        
        if format_style == "interview":
            question_text = story.get('question', '')
            if question_text:
                clean_question = clean_text(question_text)
                escaped_question = clean_question.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                rtf_parts.append(f"\\pard\\li720 \\b\\i {escaped_question}\\b0\\i0\\par\\par")
        
        answer_text = story.get('answer_text', '')
        if answer_text:
            clean_answer = clean_text(answer_text)
            paragraphs = clean_answer.split('\n')
            for para in paragraphs:
                if para.strip():
                    escaped_para = para.strip().replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                    rtf_parts.append(f"\\pard\\fi720 {escaped_para}\\par")
            rtf_parts.append("\\par")
        
        # Add image placeholders
        if include_images and story.get('images'):
            for img in story.get('images', []):
                rtf_parts.append("\\pard\\qc [Image: ")
                if img.get('caption'):
                    escaped_caption = clean_text(img['caption']).replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                    rtf_parts.append(escaped_caption)
                rtf_parts.append("]\\par\\par")
        
        rtf_parts.append("\\par")
    
    # RTF Footer
    rtf_parts.append("}")
    
    rtf_content = ''.join(rtf_parts)
    
    # Convert to bytes
    return rtf_content.encode('utf-8')

def main():
    st.set_page_config(
        page_title="Biography Publisher",
        page_icon="📚",
        layout="wide"
    )
    
    st.title("📚 Biography Publisher")
    st.markdown("Convert your stories into professional book formats")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("Book Settings")
        
        # Book metadata
        title = st.text_input("Book Title", "My Story Collection")
        author = st.text_input("Author Name", "Anonymous")
        
        # Format selection with new buttons
        st.subheader("Export Formats")
        
        # Create columns for format checkboxes
        col1, col2 = st.columns(2)
        with col1:
            export_docx = st.checkbox("📄 DOCX", value=True, help="Microsoft Word format")
            export_html = st.checkbox("🌐 HTML", value=True, help="Web page format")
            export_epub = st.checkbox("📱 EPUB", value=False, help="E-book format for readers")
        with col2:
            export_pdf = st.checkbox("📑 PDF/X", value=False, help="Print-ready PDF format")
            export_rtf = st.checkbox("📝 RTF", value=False, help="Rich Text Format")
        
        # Cover options
        st.subheader("Cover Options")
        cover_choice = st.radio(
            "Cover Type",
            ["simple", "uploaded"],
            format_func=lambda x: "Simple Text Cover" if x == "simple" else "Upload Cover Image",
            help="Choose between text-only cover or uploaded image"
        )
        
        cover_image = None
        if cover_choice == "uploaded":
            uploaded_cover = st.file_uploader("Upload Cover Image", type=['png', 'jpg', 'jpeg'])
            if uploaded_cover:
                cover_image = uploaded_cover.read()
                st.image(cover_image, width=200, caption="Cover Preview")
        
        # Formatting options
        st.subheader("Formatting")
        format_style = st.selectbox(
            "Story Format",
            ["interview", "narrative"],
            format_func=lambda x: "Interview (Q&A)" if x == "interview" else "Narrative",
            help="Choose how stories are formatted"
        )
        
        include_toc = st.checkbox("Include Table of Contents", value=True)
        include_images = st.checkbox("Include Images", value=True)
        
        # Show selected formats summary
        selected_formats = []
        if export_docx: selected_formats.append("DOCX")
        if export_html: selected_formats.append("HTML")
        if export_epub: selected_formats.append("EPUB")
        if export_pdf: selected_formats.append("PDF")
        if export_rtf: selected_formats.append("RTF")
        
        if selected_formats:
            st.info(f"📥 Will generate: {', '.join(selected_formats)}")
    
    # Main content area - Story input
    st.header("Add Your Stories")
    
    # Initialize session state for stories
    if 'stories' not in st.session_state:
        st.session_state.stories = []
    
    # Session management
    col1, col2 = st.columns([3, 1])
    with col1:
        session_title = st.text_input("Session Title", "Session 1", help="Group related stories together")
    with col2:
        if st.button("➕ New Session", use_container_width=True):
            st.session_state.stories.append({
                'session_title': session_title,
                'question': '',
                'answer_text': '',
                'images': []
            })
            st.rerun()
    
    # Display existing stories
    if not st.session_state.stories:
        st.info("👆 Click 'New Session' to start adding your first story")
    
    for i, story in enumerate(st.session_state.stories):
        with st.expander(f"📖 Story {i+1}: {story.get('session_title', 'Untitled')}", expanded=i==len(st.session_state.stories)-1):
            col1, col2 = st.columns([3, 1])
            with col1:
                story['session_title'] = st.text_input("Session Title", story['session_title'], key=f"session_{i}")
            with col2:
                if st.button(f"🗑️ Delete Story", key=f"del_{i}", use_container_width=True):
                    st.session_state.stories.pop(i)
                    st.rerun()
            
            story['question'] = st.text_area("Question/Prompt", story.get('question', ''), key=f"q_{i}", height=100)
            story['answer_text'] = st.text_area("Answer/Story", story.get('answer_text', ''), height=200, key=f"a_{i}")
            
            # Image upload for each story
            uploaded_files = st.file_uploader(
                f"Upload Images for Story {i+1}",
                type=['png', 'jpg', 'jpeg'],
                accept_multiple_files=True,
                key=f"img_{i}",
                help="Add images to accompany this story"
            )
            
            if uploaded_files:
                if 'images' not in story:
                    story['images'] = []
                
                for uploaded_file in uploaded_files:
                    img_bytes = uploaded_file.read()
                    img_base64 = base64.b64encode(img_bytes).decode()
                    
                    # Get caption
                    caption = st.text_input(f"Caption for {uploaded_file.name}", key=f"cap_{i}_{uploaded_file.name}")
                    
                    story['images'].append({
                        'base64': img_base64,
                        'caption': caption,
                        'filename': uploaded_file.name
                    })
            
            # Display existing images
            if story.get('images'):
                st.write("📸 Uploaded Images:")
                cols = st.columns(min(3, len(story['images'])))
                for j, img in enumerate(story['images']):
                    with cols[j % 3]:
                        st.image(base64.b64decode(img['base64']), width=150, caption=img.get('caption', 'No caption'))
                        if st.button(f"Remove Image", key=f"rm_{i}_{j}"):
                            story['images'].pop(j)
                            st.rerun()
    
    # Add new story button
    if st.button("➕ Add Another Story", use_container_width=True):
        st.session_state.stories.append({
            'session_title': session_title,
            'question': '',
            'answer_text': '',
            'images': []
        })
        st.rerun()
    
    # Generate button
    st.markdown("---")
    if st.button("📖 Generate Book", type="primary", use_container_width=True):
        if not st.session_state.stories:
            st.error("Please add at least one story before generating the book.")
        else:
            with st.spinner("🔄 Generating your book... This may take a moment."):
                generated_files = {}
                progress_bar = st.progress(0)
                
                # Generate each selected format
                formats_to_generate = []
                if export_docx: formats_to_generate.append(('docx', '📄 DOCX'))
                if export_html: formats_to_generate.append(('html', '🌐 HTML'))
                if export_epub: formats_to_generate.append(('epub', '📱 EPUB'))
                if export_pdf: formats_to_generate.append(('pdf', '📑 PDF'))
                if export_rtf: formats_to_generate.append(('rtf', '📝 RTF'))
                
                for idx, (fmt, label) in enumerate(formats_to_generate):
                    progress_bar.progress((idx) / len(formats_to_generate), text=f"Generating {label}...")
                    
                    if fmt == 'docx':
                        data = generate_docx(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                    elif fmt == 'html':
                        html_str = generate_html(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                        data = html_str.encode('utf-8')
                    elif fmt == 'epub':
                        data = generate_epub(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                    elif fmt == 'pdf':
                        data = generate_pdf(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                    elif fmt == 'rtf':
                        data = generate_rtf(title, author, st.session_state.stories, format_style, include_toc, include_images, cover_image, cover_choice)
                    
                    generated_files[fmt] = data
                
                progress_bar.progress(1.0, text="✅ Generation complete!")
                progress_bar.empty()
                
                # Show celebration
                show_celebration()
                
                # Display download buttons
                st.subheader("📥 Download Your Book")
                
                if generated_files:
                    # Create columns for download buttons
                    cols = st.columns(len(generated_files))
                    
                    mime_types = {
                        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'html': 'text/html',
                        'epub': 'application/epub+zip',
                        'pdf': 'application/pdf',
                        'rtf': 'application/rtf'
                    }
                    
                    button_icons = {
                        'docx': '📄',
                        'html': '🌐',
                        'epub': '📱',
                        'pdf': '📑',
                        'rtf': '📝'
                    }
                    
                    for idx, (fmt, data) in enumerate(generated_files.items()):
                        with cols[idx]:
                            file_name = f"{title.replace(' ', '_')}.{fmt}"
                            st.download_button(
                                label=f"{button_icons.get(fmt, '📥')} Download {fmt.upper()}",
                                data=data,
                                file_name=file_name,
                                mime=mime_types.get(fmt, 'application/octet-stream'),
                                use_container_width=True
                            )
                else:
                    st.warning("No formats selected. Please check at least one format in the sidebar.")

if __name__ == "__main__":
    main()
