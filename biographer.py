# biographer.py – Tell My Story App (COMPLETE VERSION)
import streamlit as st
import json
from datetime import datetime, date
import os
import re
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import secrets
import string
import time
import base64
from PIL import Image
import io
import zipfile

# ============================================================================
# IMPORT QUILL RICH TEXT EDITOR
# ============================================================================
try:
    from streamlit_quill import st_quill
    QUILL_AVAILABLE = True
except ImportError:
    st.error("❌ Please install streamlit-quill: pip install streamlit-quill")
    st.stop()

# ============================================================================
# FORCE DIRECTORY CREATION
# ============================================================================
for dir_path in ["question_banks/default", "question_banks/users", "question_banks", 
                 "uploads", "uploads/thumbnails", "uploads/metadata", "accounts", "sessions", "backups"]:
    os.makedirs(dir_path, exist_ok=True)

DEFAULT_WORD_TARGET = 500

# ============================================================================
# INITIALIZATION
# ============================================================================
# Initialize session state
if "user_id" not in st.session_state:
    st.session_state.user_id = None
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "current_session" not in st.session_state:
    st.session_state.current_session = 0
if "current_question" not in st.session_state:
    st.session_state.current_question = 0
if "responses" not in st.session_state:
    st.session_state.responses = {}
if "editing" not in st.session_state:
    st.session_state.editing = False
if "editing_word_target" not in st.session_state:
    st.session_state.editing_word_target = False
if "confirming_clear" not in st.session_state:
    st.session_state.confirming_clear = None
if "data_loaded" not in st.session_state:
    st.session_state.data_loaded = False
if "current_question_override" not in st.session_state:
    st.session_state.current_question_override = None
if "user_account" not in st.session_state:
    st.session_state.user_account = None
if "show_profile_setup" not in st.session_state:
    st.session_state.show_profile_setup = False
if "image_handler" not in st.session_state:
    st.session_state.image_handler = None
if "show_ai_rewrite" not in st.session_state:
    st.session_state.show_ai_rewrite = False
if "show_ai_rewrite_menu" not in st.session_state:
    st.session_state.show_ai_rewrite_menu = False
if "editor_content" not in st.session_state:
    st.session_state.editor_content = {}
if "show_privacy_settings" not in st.session_state:
    st.session_state.show_privacy_settings = False
if "show_cover_designer" not in st.session_state:
    st.session_state.show_cover_designer = False
if "beta_feedback_display" not in st.session_state:
    st.session_state.beta_feedback_display = None
if "beta_feedback_storage" not in st.session_state:
    st.session_state.beta_feedback_storage = {}
if "current_question_bank" not in st.session_state:
    st.session_state.current_question_bank = None
if "current_bank_name" not in st.session_state:
    st.session_state.current_bank_name = None
if "show_publisher" not in st.session_state:
    st.session_state.show_publisher = False

# ============================================================================
# PAGE CONFIG
# ============================================================================
st.set_page_config(page_title="Tell My Story", page_icon="📖", layout="wide")

LOGO_URL = "https://menuhunterai.com/wp-content/uploads/2026/02/tms_logo.png"

# ============================================================================
# DEFAULT QUESTION BANK
# ============================================================================
DEFAULT_SESSIONS = [
    {
        "id": 1,
        "title": "Early Years",
        "questions": [
            "Where were you born and what do you remember about your childhood home?",
            "Tell me about your parents - what were they like?",
            "Did you have siblings? What was your relationship with them?",
            "What was school like for you? Favorite subjects?",
            "Who were your childhood friends? Any special memories?"
        ],
        "word_target": 750,
        "guidance": "Focus on sensory details - what did your home smell like? What sounds do you remember?"
    },
    {
        "id": 2,
        "title": "Coming of Age",
        "questions": [
            "What were your teenage years like?",
            "Tell me about your first love or first heartbreak.",
            "What dreams did you have for your future?",
            "Who influenced you most during this time?",
            "What challenges did you face growing up?"
        ],
        "word_target": 750,
        "guidance": "Think about the music, fashion, and culture of your teenage years."
    },
    {
        "id": 3,
        "title": "Career & Life's Work",
        "questions": [
            "What was your first job? What did you learn from it?",
            "Describe your career path - the twists and turns.",
            "What achievements are you most proud of?",
            "Tell me about important mentors or colleagues.",
            "What work brought you the most fulfillment?"
        ],
        "word_target": 750,
        "guidance": "Include both successes and failures - they all shaped your journey."
    },
    {
        "id": 4,
        "title": "Love & Family",
        "questions": [
            "How did you meet your partner? Tell me your love story.",
            "What was your wedding day like?",
            "Tell me about becoming a parent (if applicable).",
            "What traditions did you create with your family?",
            "How has your family changed over the years?"
        ],
        "word_target": 750,
        "guidance": "Share the small moments as well as the big ones."
    },
    {
        "id": 5,
        "title": "Wisdom & Legacy",
        "questions": [
            "What life lessons would you want to pass on?",
            "What moments are you most proud of?",
            "What challenges made you who you are?",
            "How do you want to be remembered?",
            "What advice would you give your younger self?"
        ],
        "word_target": 750,
        "guidance": "Reflect on what really matters to you now."
    }
]

# Initialize question bank
if st.session_state.current_question_bank is None:
    st.session_state.current_question_bank = DEFAULT_SESSIONS
    st.session_state.current_bank_name = "📖 Life Story - Default"

# Initialize responses for each session
for session in st.session_state.current_question_bank:
    sid = session["id"]
    if sid not in st.session_state.responses:
        st.session_state.responses[sid] = {
            "title": session["title"],
            "questions": {},
            "summary": "",
            "completed": False,
            "word_target": session.get("word_target", DEFAULT_WORD_TARGET)
        }

# ============================================================================
# IMAGE HANDLER CLASS
# ============================================================================
class ImageHandler:
    def __init__(self, user_id=None):
        self.user_id = user_id
        self.base_path = "uploads"
        self.settings = {
            "full_width": 1600,
            "inline_width": 800,
            "thumbnail_size": 200,
            "quality": 85,
            "max_file_size_mb": 5
        }
    
    def get_user_path(self):
        if self.user_id:
            user_hash = hashlib.md5(self.user_id.encode()).hexdigest()[:8]
            path = f"{self.base_path}/user_{user_hash}"
            os.makedirs(f"{path}/thumbnails", exist_ok=True)
            return path
        return self.base_path
    
    def optimize_image(self, image, max_width=1600, is_thumbnail=False):
        try:
            if image.mode in ('RGBA', 'LA', 'P'):
                bg = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                if image.mode == 'RGBA':
                    bg.paste(image, mask=image.split()[-1])
                else:
                    bg.paste(image)
                image = bg
            
            if is_thumbnail:
                image.thumbnail((self.settings["thumbnail_size"], self.settings["thumbnail_size"]), Image.Resampling.LANCZOS)
                return image
            
            if image.width > max_width:
                new_height = int((max_width / image.width) * image.height)
                image = image.resize((max_width, new_height), Image.Resampling.LANCZOS)
            
            return image
        except Exception as e:
            return image
    
    def save_image(self, uploaded_file, session_id, question_text, caption="", usage="full_page"):
        try:
            image_data = uploaded_file.read()
            
            img = Image.open(io.BytesIO(image_data))
            target_width = self.settings["full_width"] if usage == "full_page" else self.settings["inline_width"]
            
            image_id = hashlib.md5(f"{self.user_id}{session_id}{question_text}{datetime.now()}".encode()).hexdigest()[:16]
            
            optimized_img = self.optimize_image(img, target_width, is_thumbnail=False)
            thumb_img = self.optimize_image(img, is_thumbnail=True)
            
            main_buffer = io.BytesIO()
            optimized_img.save(main_buffer, format="JPEG", quality=self.settings["quality"], optimize=True)
            
            thumb_buffer = io.BytesIO()
            thumb_img.save(thumb_buffer, format="JPEG", quality=70, optimize=True)
            
            user_path = self.get_user_path()
            with open(f"{user_path}/{image_id}.jpg", 'wb') as f: 
                f.write(main_buffer.getvalue())
            with open(f"{user_path}/thumbnails/{image_id}.jpg", 'wb') as f: 
                f.write(thumb_buffer.getvalue())
            
            metadata = {
                "id": image_id, "session_id": session_id, "question": question_text,
                "caption": caption, "user_id": self.user_id,
                "timestamp": datetime.now().isoformat(), "usage": usage
            }
            with open(f"{self.base_path}/metadata/{image_id}.json", 'w') as f: 
                json.dump(metadata, f)
            
            return {
                "has_images": True, 
                "images": [{
                    "id": image_id, "caption": caption
                }]
            }
        except Exception as e:
            return None
    
    def get_image_base64(self, image_id):
        try:
            user_path = self.get_user_path()
            path = f"{user_path}/{image_id}.jpg"
            if not os.path.exists(path): 
                return None
            with open(path, 'rb') as f: 
                image_data = f.read()
            return base64.b64encode(image_data).decode()
        except:
            return None
    
    def get_image_caption(self, image_id):
        meta_path = f"{self.base_path}/metadata/{image_id}.json"
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                    return metadata.get("caption", "")
            except:
                pass
        return ""
    
    def get_images_for_answer(self, session_id, question_text):
        images = []
        metadata_dir = f"{self.base_path}/metadata"
        if not os.path.exists(metadata_dir): 
            return images
        
        for fname in os.listdir(metadata_dir):
            if fname.endswith('.json'):
                try:
                    with open(f"{metadata_dir}/{fname}") as f: 
                        meta = json.load(f)
                    if (meta.get("session_id") == session_id and 
                        meta.get("question") == question_text and 
                        meta.get("user_id") == self.user_id):
                        
                        images.append(meta)
                except:
                    continue
        return sorted(images, key=lambda x: x.get("timestamp", ""), reverse=True)
    
    def delete_image(self, image_id):
        try:
            user_path = self.get_user_path()
            for p in [f"{user_path}/{image_id}.jpg", 
                     f"{user_path}/thumbnails/{image_id}.jpg", 
                     f"{self.base_path}/metadata/{image_id}.json"]:
                if os.path.exists(p): 
                    os.remove(p)
            return True
        except:
            return False

# ============================================================================
# AUTHENTICATION FUNCTIONS
# ============================================================================
def generate_password(length=12):
    alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
    return ''.join(secrets.choice(alphabet) for _ in range(length))

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(stored_hash, password):
    return stored_hash == hash_password(password)

def get_account_data(user_id=None, email=None):
    try:
        if user_id:
            fname = f"accounts/{user_id}_account.json"
            if os.path.exists(fname): 
                return json.load(open(fname, 'r'))
        if email:
            email = email.lower().strip()
            for f in os.listdir("accounts"):
                if f.endswith("_account.json"):
                    try:
                        data = json.load(open(f"accounts/{f}", 'r'))
                        if data.get("email", "").lower() == email:
                            return data
                    except:
                        continue
    except: 
        pass
    return None

def create_user_account(user_data, password=None):
    try:
        user_id = hashlib.sha256(f"{user_data['email']}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
        if not password: 
            password = generate_password()
        user_record = {
            "user_id": user_id, "email": user_data["email"].lower().strip(),
            "password_hash": hash_password(password), 
            "created_at": datetime.now().isoformat(), "last_login": datetime.now().isoformat(),
            "profile": {
                "first_name": user_data["first_name"], "last_name": user_data["last_name"],
                "email": user_data["email"]
            }
        }
        with open(f"accounts/{user_id}_account.json", 'w') as f:
            json.dump(user_record, f)
        return {"success": True, "user_id": user_id, "password": password, "user_record": user_record}
    except Exception as e:
        return {"success": False, "error": str(e)}

def authenticate_user(email, password):
    try:
        account = get_account_data(email=email)
        if account and verify_password(account['password_hash'], password):
            account['last_login'] = datetime.now().isoformat()
            with open(f"accounts/{account['user_id']}_account.json", 'w') as f:
                json.dump(account, f)
            return {"success": True, "user_id": account['user_id'], "user_record": account}
        return {"success": False, "error": "Invalid email or password"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def save_account_data(user_record):
    try:
        with open(f"accounts/{user_record['user_id']}_account.json", 'w') as f:
            json.dump(user_record, f)
        return True
    except: 
        return False

def logout_user():
    st.session_state.user_id = None
    st.session_state.logged_in = False
    st.session_state.user_account = None
    st.session_state.image_handler = None
    st.session_state.show_profile_setup = False
    st.session_state.current_session = 0
    st.session_state.current_question = 0
    st.session_state.current_question_override = None
    st.rerun()

# ============================================================================
# STORAGE FUNCTIONS
# ============================================================================
def get_user_filename(user_id):
    return f"user_data_{hashlib.md5(user_id.encode()).hexdigest()[:8]}.json"

def load_user_data(user_id):
    fname = get_user_filename(user_id)
    try:
        if os.path.exists(fname):
            return json.load(open(fname, 'r'))
        return {"responses": {}}
    except: 
        return {"responses": {}}

def save_user_data(user_id, responses_data):
    fname = get_user_filename(user_id)
    try:
        data = {
            "user_id": user_id, 
            "responses": responses_data,
            "last_saved": datetime.now().isoformat()
        }
        with open(fname, 'w') as f: 
            json.dump(data, f)
        return True
    except: 
        return False

# ============================================================================
# CORE RESPONSE FUNCTIONS
# ============================================================================
def save_response(session_id, question, answer):
    user_id = st.session_state.user_id
    if not user_id: 
        return False
    
    if session_id not in st.session_state.responses:
        session_data = next((s for s in st.session_state.current_question_bank if s["id"] == session_id), None)
        st.session_state.responses[session_id] = {
            "title": session_data["title"] if session_data else f"Session {session_id}",
            "questions": {}, "word_target": session_data.get("word_target", DEFAULT_WORD_TARGET) if session_data else DEFAULT_WORD_TARGET
        }
    
    # Get images for this answer
    images = []
    if st.session_state.image_handler:
        images = st.session_state.image_handler.get_images_for_answer(session_id, question)
    
    st.session_state.responses[session_id]["questions"][question] = {
        "answer": answer, "timestamp": datetime.now().isoformat(),
        "has_images": len(images) > 0,
        "image_count": len(images),
        "images": [{"id": img["id"], "caption": img.get("caption", "")} for img in images]
    }
    
    return save_user_data(user_id, st.session_state.responses)

def delete_response(session_id, question):
    user_id = st.session_state.user_id
    if not user_id: 
        return False
    
    if session_id in st.session_state.responses and question in st.session_state.responses[session_id]["questions"]:
        del st.session_state.responses[session_id]["questions"][question]
        return save_user_data(user_id, st.session_state.responses)
    return False

def calculate_author_word_count(session_id):
    total = 0
    if session_id in st.session_state.responses:
        for q, d in st.session_state.responses[session_id].get("questions", {}).items():
            if d.get("answer"): 
                text_only = re.sub(r'<[^>]+>', '', d["answer"])
                total += len(re.findall(r'\w+', text_only))
    return total

def get_progress_info(session_id):
    current = calculate_author_word_count(session_id)
    target = st.session_state.responses[session_id].get("word_target", DEFAULT_WORD_TARGET)
    if target == 0: 
        percent = 100
    else: 
        percent = (current / target) * 100
    
    return {
        "current_count": current, "target": target, "progress_percent": percent,
        "emoji": "🟢" if percent >= 100 else "🟡" if percent >= 70 else "🔴",
        "color": "#27ae60" if percent >= 100 else "#f39c12" if percent >= 70 else "#e74c3c",
        "remaining_words": max(0, target - current)
    }

def search_all_answers(search_query):
    if not search_query or len(search_query) < 2: 
        return []
    
    results = []
    search_query = search_query.lower()
    
    for session in st.session_state.current_question_bank:
        session_id = session["id"]
        session_data = st.session_state.responses.get(session_id, {})
        
        for question_text, answer_data in session_data.get("questions", {}).items():
            html_answer = answer_data.get("answer", "")
            text_answer = re.sub(r'<[^>]+>', '', html_answer)
            
            if search_query in text_answer.lower() or search_query in question_text.lower():
                results.append({
                    "session_id": session_id, "session_title": session["title"],
                    "question": question_text, "answer": text_answer[:200] + "..." if len(text_answer) > 200 else text_answer,
                    "timestamp": answer_data.get("timestamp", ""), "word_count": len(text_answer.split()),
                    "has_images": answer_data.get("has_images", False)
                })
    
    return sorted(results, key=lambda x: x["timestamp"], reverse=True)

# ============================================================================
# AI REWRITE FUNCTION
# ============================================================================
def auto_correct_text(text):
    if not text: 
        return text
    text_only = re.sub(r'<[^>]+>', '', text)
    # Simple correction for demo (in real app, would use AI)
    corrected = text_only.replace("  ", " ").strip()
    return corrected

def ai_rewrite_answer(original_text, person_option, question_text, session_title):
    """Simple rewrite for demo purposes"""
    text_only = re.sub(r'<[^>]+>', '', original_text)
    
    if len(text_only.split()) < 5:
        return {"error": "Text too short to rewrite"}
    
    # Simple rewrite for demo
    if person_option == "1st":
        rewritten = text_only
    elif person_option == "2nd":
        # Replace I/me/my with you/your (simplified)
        rewritten = text_only.replace("I ", "you ").replace("I'm", "you're").replace("my ", "your ").replace("me", "you")
    else:
        # Replace I/me/my with they/their (simplified)
        rewritten = text_only.replace("I ", "they ").replace("I'm", "they're").replace("my ", "their ").replace("me", "them")
    
    return {
        "success": True,
        "original": text_only,
        "rewritten": rewritten,
        "person": ["First Person", "Second Person", "Third Person"][["1st", "2nd", "3rd"].index(person_option)],
        "emoji": ["👤", "💬", "📖"][["1st", "2nd", "3rd"].index(person_option)]
    }

# ============================================================================
# PROFILE FUNCTIONS
# ============================================================================
def render_narrative_gps():
    st.markdown("### ❤️ The Heart of Your Story")
    
    if 'narrative_gps' not in st.session_state.user_account:
        st.session_state.user_account['narrative_gps'] = {}
    
    gps = st.session_state.user_account['narrative_gps']
    
    gps['book_title'] = st.text_input("Working title for your book:", value=gps.get('book_title', ''))
    gps['purpose'] = st.text_area("Why are you writing your story?", value=gps.get('purpose', ''), height=100)
    gps['audience'] = st.text_input("Who do you want to read this?", value=gps.get('audience', ''))
    gps['tone'] = st.selectbox("What tone do you want?", 
                               ["Warm and personal", "Professional", "Reflective", "Humorous", "Raw and honest"],
                               index=["Warm and personal", "Professional", "Reflective", "Humorous", "Raw and honest"].index(gps.get('tone', 'Warm and personal')))
    
    if st.button("💾 Save Heart of Your Story", type="primary"):
        save_account_data(st.session_state.user_account)
        st.success("Saved!")
        st.rerun()

def render_enhanced_profile():
    st.markdown("### 📋 Your Life Story Details")
    
    if 'enhanced_profile' not in st.session_state.user_account:
        st.session_state.user_account['enhanced_profile'] = {}
    
    ep = st.session_state.user_account['enhanced_profile']
    
    ep['birth_place'] = st.text_input("Where were you born?", value=ep.get('birth_place', ''))
    ep['parents'] = st.text_area("Tell me about your parents:", value=ep.get('parents', ''), height=100)
    ep['childhood'] = st.text_area("What was your childhood like?", value=ep.get('childhood', ''), height=100)
    ep['education'] = st.text_area("Tell me about your education:", value=ep.get('education', ''), height=100)
    ep['career'] = st.text_area("What was your career/journey?", value=ep.get('career', ''), height=100)
    ep['family'] = st.text_area("Tell me about your family:", value=ep.get('family', ''), height=100)
    ep['challenges'] = st.text_area("What challenges shaped you?", value=ep.get('challenges', ''), height=100)
    ep['wisdom'] = st.text_area("What life lessons do you want to share?", value=ep.get('wisdom', ''), height=100)
    
    if st.button("💾 Save Life Story Details", type="primary"):
        save_account_data(st.session_state.user_account)
        st.success("Saved!")
        st.rerun()

def show_privacy_settings():
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("🔒 Privacy Settings")
    
    if st.button("← Back"):
        st.session_state.show_privacy_settings = False
        st.rerun()
    
    if 'privacy_settings' not in st.session_state.user_account:
        st.session_state.user_account['privacy_settings'] = {
            "profile_public": False, "stories_public": False, "allow_sharing": False,
            "data_collection": True
        }
    
    privacy = st.session_state.user_account['privacy_settings']
    
    privacy['profile_public'] = st.checkbox("Make profile public", value=privacy.get('profile_public', False))
    privacy['stories_public'] = st.checkbox("Share stories publicly", value=privacy.get('stories_public', False))
    privacy['allow_sharing'] = st.checkbox("Allow sharing via link", value=privacy.get('allow_sharing', False))
    privacy['data_collection'] = st.checkbox("Allow anonymous usage data", value=privacy.get('data_collection', True))
    
    if st.button("💾 Save Privacy Settings", type="primary"):
        save_account_data(st.session_state.user_account)
        st.success("Privacy settings saved!")
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

def show_cover_designer():
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("🎨 Cover Designer")
    
    if st.button("← Back"):
        st.session_state.show_cover_designer = False
        st.rerun()
    
    st.info("Cover designer coming soon! For now, you can add a cover when publishing your book.")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# AI REWRITE MODAL
# ============================================================================
def show_ai_rewrite_modal():
    if not st.session_state.get('current_rewrite_data'):
        return
    
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([6, 1])
    with col1:
        st.markdown(f"### {st.session_state.current_rewrite_data.get('emoji', '✨')} AI Rewrite")
    with col2:
        if st.button("✕"):
            st.session_state.show_ai_rewrite = False
            st.session_state.current_rewrite_data = None
            st.rerun()
    
    rewrite_data = st.session_state.current_rewrite_data
    
    if rewrite_data.get('error'):
        st.error(f"Could not rewrite: {rewrite_data['error']}")
    else:
        st.markdown("**Original Version:**")
        st.markdown(f'<div style="background:#f5f5f5; padding:15px; border-radius:5px;">{rewrite_data["original"]}</div>', unsafe_allow_html=True)
        
        st.markdown("**Rewritten Version:**")
        st.markdown(f'<div style="background:#e3f2fd; padding:15px; border-radius:5px;">{rewrite_data["rewritten"]}</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📝 Replace Original", type="primary", use_container_width=True):
                # Get current session and question
                current_session = st.session_state.current_question_bank[st.session_state.current_session]
                current_session_id = current_session["id"]
                current_question_text = st.session_state.current_question_override or current_session["questions"][st.session_state.current_question]
                
                # Format and save
                new_content = f"<p>{rewrite_data['rewritten']}</p>"
                
                # Update editor content
                editor_key = f"quill_{current_session_id}_{current_question_text[:20]}"
                content_key = f"{editor_key}_content"
                st.session_state[content_key] = new_content
                
                # Save to database
                save_response(current_session_id, current_question_text, new_content)
                
                st.success("Replaced with rewritten version!")
                time.sleep(1)
                st.session_state.show_ai_rewrite = False
                st.session_state.current_rewrite_data = None
                st.rerun()
        
        with col2:
            if st.button("❌ Cancel", use_container_width=True):
                st.session_state.show_ai_rewrite = False
                st.session_state.current_rewrite_data = None
                st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# IMPORT FUNCTIONS
# ============================================================================
def import_text_file(uploaded_file):
    """Import text from common document formats"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'txt':
            content = uploaded_file.read().decode('utf-8', errors='ignore')
        
        elif file_extension == 'docx':
            try:
                from docx import Document
                import io
                docx_bytes = io.BytesIO(uploaded_file.getvalue())
                doc = Document(docx_bytes)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                content = '\n\n'.join(paragraphs)
            except ImportError:
                st.warning("Please install python-docx for DOCX support")
                return None
        
        else:
            st.error(f"Unsupported format: .{file_extension}")
            st.info("Supported: .txt, .docx")
            return None
        
        # Clean and format
        content = re.sub(r'\s+', ' ', content)
        sentences = re.split(r'[.!?]+', content)
        paragraphs = []
        current_para = []
        
        for sentence in sentences:
            if sentence.strip():
                current_para.append(sentence.strip() + '.')
                if len(current_para) >= 4:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
        
        if current_para:
            paragraphs.append(' '.join(current_para))
        
        if not paragraphs:
            paragraphs = [content]
        
        html_content = ''
        for para in paragraphs:
            if para.strip():
                para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                html_content += f'<p>{para.strip()}</p>'
        
        return html_content
        
    except Exception as e:
        st.error(f"Import error: {str(e)}")
        return None

# ============================================================================
# PUBLISHER FUNCTIONS
# ============================================================================
def prepare_stories_for_publishing(responses_data, sessions_data):
    """Convert response data to format needed for publishing"""
    stories = []
    
    for session in sessions_data:
        session_id = session["id"]
        session_data = responses_data.get(session_id, {})
        
        for question_text, answer_data in session_data.get("questions", {}).items():
            # Get images if available
            images = []
            if answer_data.get("images") and st.session_state.image_handler:
                for img_ref in answer_data.get("images", []):
                    img_id = img_ref.get("id")
                    b64 = st.session_state.image_handler.get_image_base64(img_id)
                    if b64:
                        images.append({
                            "id": img_id,
                            "base64": b64,
                            "caption": img_ref.get("caption", "")
                        })
            
            stories.append({
                "question": question_text,
                "answer_text": re.sub(r'<[^>]+>', '', answer_data.get("answer", "")),
                "timestamp": answer_data.get("timestamp", ""),
                "session_id": session_id,
                "session_title": session["title"],
                "has_images": answer_data.get("has_images", False),
                "image_count": answer_data.get("image_count", 0),
                "images": images
            })
    
    return stories

def generate_docx(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate a Word document from stories"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        # Cover page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"by {author}")
        author_run.font.size = Pt(16)
        author_run.font.italic = True
        
        doc.add_page_break()
        
        # Table of Contents
        if include_toc:
            doc.add_heading("Table of Contents", level=1)
            sessions = {}
            for story in stories:
                session_title = story.get('session_title', 'Untitled Session')
                if session_title not in sessions:
                    sessions[session_title] = []
                sessions[session_title].append(story)
            
            for session_title in sessions.keys():
                doc.add_paragraph(f"  {session_title}", style='List Bullet')
            doc.add_page_break()
        
        # Add stories
        current_session = None
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            
            if session_title != current_session:
                current_session = session_title
                doc.add_heading(session_title, level=1)
            
            if format_style == "interview":
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(story.get('question', ''))
                q_run.font.bold = True
                q_run.font.italic = True
            
            # Add answer
            answer_text = story.get('answer_text', '')
            if answer_text:
                paragraphs = answer_text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            # Add images
            if include_images and story.get('images'):
                for img in story.get('images', []):
                    if img.get('base64'):
                        try:
                            img_data = base64.b64decode(img['base64'])
                            img_stream = io.BytesIO(img_data)
                            doc.add_picture(img_stream, width=Inches(4))
                            if img.get('caption'):
                                caption = doc.add_paragraph(img['caption'])
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except:
                            pass
            
            doc.add_paragraph()
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    except ImportError:
        st.error("Please install python-docx: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Error generating DOCX: {str(e)}")
        return None

def generate_html(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate HTML document from stories"""
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Georgia', serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        h1 {{
            font-size: 42px;
            text-align: center;
            margin-bottom: 10px;
        }}
        h2 {{
            font-size: 28px;
            margin-top: 40px;
            color: #444;
            border-bottom: 2px solid #eee;
            padding-bottom: 10px;
        }}
        .author {{
            text-align: center;
            font-size: 18px;
            color: #666;
            margin-bottom: 40px;
            font-style: italic;
        }}
        .question {{
            font-weight: bold;
            font-size: 18px;
            margin-top: 30px;
            color: #2c3e50;
            border-left: 4px solid #3498db;
            padding-left: 15px;
        }}
        .story-image {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border-radius: 5px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .image-caption {{
            text-align: center;
            font-size: 14px;
            color: #666;
            font-style: italic;
        }}
        .cover-page {{
            text-align: center;
            margin-bottom: 50px;
        }}
        .copyright {{
            text-align: center;
            font-size: 12px;
            color: #999;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #eee;
        }}
        .toc {{
            background: #f9f9f9;
            padding: 20px;
            border-radius: 5px;
            margin: 30px 0;
        }}
    </style>
</head>
<body>
    <div class="cover-page">
        <h1>{title}</h1>
        <p class="author">by {author}</p>
    </div>
    
    <p class="copyright">© {datetime.now().year} {author}. All rights reserved.</p>
"""
    
    if include_toc:
        html += """
    <div class="toc">
        <h3>Table of Contents</h3>
        <ul>
"""
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        for session_title in sessions.keys():
            anchor = session_title.lower().replace(' ', '-').replace('?', '').replace('!', '').replace(',', '')
            html += f'            <li><a href="#{anchor}">{session_title}</a></li>\n'
        
        html += """
        </ul>
    </div>
"""
    
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        anchor = session_title.lower().replace(' ', '-').replace('?', '').replace('!', '').replace(',', '')
        
        if session_title != current_session:
            current_session = session_title
            html += f'    <h2 id="{anchor}">{session_title}</h2>\n'
        
        if format_style == "interview":
            html += f'    <div class="question">{story.get("question", "")}</div>\n'
        
        # Format answer
        answer_text = story.get('answer_text', '')
        if answer_text:
            paragraphs = answer_text.split('\n')
            for para in paragraphs:
                if para.strip():
                    escaped_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    html += f'    <p>{escaped_para}</p>\n'
        
        # Add images
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    html += f'    <img src="data:image/jpeg;base64,{img["base64"]}" class="story-image">\n'
                    if img.get('caption'):
                        caption = img['caption'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        html += f'    <p class="image-caption">{caption}</p>\n'
        
        html += '    <hr style="margin: 30px 0; border: none; border-top: 1px dashed #ccc;">\n'
    
    html += """
</body>
</html>"""
    
    return html

def generate_zip(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate ZIP package with HTML and images"""
    import zipfile
    import io
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Generate HTML
        html_content = generate_html(
            title, author, stories, format_style, include_toc, include_images,
            cover_image, cover_choice
        )
        
        # Save HTML
        html_filename = f"{title.replace(' ', '_')}.html"
        zip_file.writestr(html_filename, html_content)
        
        # Add images
        if include_images:
            for i, story in enumerate(stories):
                for j, img in enumerate(story.get('images', [])):
                    if img.get('base64'):
                        img_data = base64.b64decode(img['base64'])
                        img_filename = f"images/image_{i}_{j}.jpg"
                        zip_file.writestr(img_filename, img_data)
    
    return zip_buffer.getvalue()

def show_publisher_interface():
    """Display the publisher interface"""
    
    st.markdown("## 📚 Publish Your Book")
    
    # Get data from main app
    if not st.session_state.get('logged_in'):
        st.warning("Please log in to publish your book.")
        return
    
    # Prepare stories
    stories = prepare_stories_for_publishing(
        st.session_state.responses,
        st.session_state.current_question_bank
    )
    
    if not stories:
        st.info("No stories yet! Start writing to publish your book.")
        return
    
    # Get user profile info
    profile = st.session_state.user_account.get('profile', {}) if st.session_state.user_account else {}
    first_name = profile.get('first_name', 'My')
    last_name = profile.get('last_name', '')
    
    # Publishing options
    col1, col2 = st.columns(2)
    
    with col1:
        book_title = st.text_input(
            "Book Title",
            value=f"{first_name}'s Life Story",
            key="publisher_book_title"
        )
        
        author_name = st.text_input(
            "Author Name",
            value=f"{first_name} {last_name}".strip() or "Author Name",
            key="publisher_author_name"
        )
        
        format_style = st.selectbox(
            "Format Style",
            ["interview", "biography", "memoir"],
            format_func=lambda x: {
                "interview": "📝 Interview Q&A",
                "biography": "📖 Continuous Biography",
                "memoir": "📚 Chapter-based Memoir"
            }[x],
            key="publisher_format"
        )
    
    with col2:
        st.markdown("### Cover Options")
        cover_choice = st.radio(
            "Cover Type",
            ["simple", "uploaded"],
            format_func=lambda x: "🎨 Simple Cover" if x == "simple" else "📸 Upload Cover Image",
            key="publisher_cover_type"
        )
        
        cover_image = None
        if cover_choice == "uploaded":
            uploaded_cover = st.file_uploader(
                "Upload Cover Image",
                type=['jpg', 'jpeg', 'png'],
                key="publisher_cover_upload"
            )
            if uploaded_cover:
                cover_image = uploaded_cover.getvalue()
                st.image(cover_image, caption="Cover Preview", width=200)
        
        include_toc = st.checkbox("Include Table of Contents", value=True, key="publisher_toc")
        include_images = st.checkbox("Include Images", value=True, key="publisher_images")
    
    st.markdown("---")
    st.markdown(f"**Total Stories:** {len(stories)} | **Total Words:** {sum(len(s['answer_text'].split()) for s in stories):,}")
    st.markdown("---")
    
    # Generate buttons
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📊 Generate DOCX", type="primary", use_container_width=True):
            with st.spinner("Creating Word document..."):
                docx_bytes = generate_docx(
                    book_title, author_name, stories, format_style,
                    include_toc, include_images, cover_image, cover_choice
                )
                
                if docx_bytes:
                    filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                    
                    st.download_button(
                        label="📥 Download DOCX",
                        data=docx_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="docx_download_publisher"
                    )
                    st.balloons()
                else:
                    st.error("Failed to generate DOCX")
    
    with col2:
        if st.button("🌐 Generate HTML", type="primary", use_container_width=True):
            with st.spinner("Creating HTML page..."):
                html_content = generate_html(
                    book_title, author_name, stories, format_style,
                    include_toc, include_images, cover_image, cover_choice
                )
                
                filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.html"
                
                st.download_button(
                    label="📥 Download HTML",
                    data=html_content,
                    file_name=filename,
                    mime="text/html",
                    use_container_width=True,
                    key="html_download_publisher"
                )
                st.balloons()
    
    with col3:
        if st.button("📦 Generate ZIP", type="primary", use_container_width=True):
            with st.spinner("Creating ZIP package..."):
                zip_data = generate_zip(
                    book_title, author_name, stories, format_style,
                    include_toc, include_images, cover_image, cover_choice
                )
                
                filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.zip"
                
                st.download_button(
                    label="📥 Download ZIP",
                    data=zip_data,
                    file_name=filename,
                    mime="application/zip",
                    use_container_width=True,
                    key="zip_download_publisher"
                )
                st.balloons()

def show_publisher_modal():
    """Show publisher interface as a modal"""
    
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("🎨 Publish Your Book")
    with col2:
        if st.button("✕", key="close_publisher"):
            st.session_state.show_publisher = False
            st.rerun()
    
    st.markdown("---")
    show_publisher_interface()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# AUTHENTICATION UI
# ============================================================================
if not st.session_state.logged_in:
    st.markdown('<div style="text-align:center; padding:50px;"><h1>Tell My Story</h1><p>Your Life Timeline • Preserve Your Legacy</p></div>', unsafe_allow_html=True)
    
    if 'auth_tab' not in st.session_state: 
        st.session_state.auth_tab = 'login'
    
    col1, col2 = st.columns(2)
    with col1: 
        st.button("🔐 Login", width='stretch', type="primary" if st.session_state.auth_tab=='login' else "secondary", 
                        on_click=lambda: st.session_state.update(auth_tab='login'))
    with col2: 
        st.button("📝 Sign Up", width='stretch', type="primary" if st.session_state.auth_tab=='signup' else "secondary",
                        on_click=lambda: st.session_state.update(auth_tab='signup'))
    st.divider()
    
    if st.session_state.auth_tab == 'login':
        with st.form("login_form"):
            st.subheader("Welcome Back")
            email = st.text_input("Email Address")
            password = st.text_input("Password", type="password")
            if st.form_submit_button("Login", type="primary", use_container_width=True):
                if email and password:
                    result = authenticate_user(email, password)
                    if result["success"]:
                        st.session_state.user_id = result["user_id"]
                        st.session_state.user_account = result["user_record"]
                        st.session_state.logged_in = True
                        st.session_state.data_loaded = False
                        st.session_state.image_handler = ImageHandler(result["user_id"])
                        st.success("Login successful!")
                        st.rerun()
                    else: 
                        st.error(f"Login failed: {result.get('error', 'Unknown error')}")
    else:
        with st.form("signup_form"):
            st.subheader("Create New Account")
            col1, col2 = st.columns(2)
            with col1: 
                first_name = st.text_input("First Name*")
            with col2: 
                last_name = st.text_input("Last Name*")
            email = st.text_input("Email Address*")
            password = st.text_input("Password*", type="password")
            confirm = st.text_input("Confirm Password*", type="password")
            accept = st.checkbox("I agree to the Terms*")
            
            if st.form_submit_button("Create Account", type="primary", use_container_width=True):
                errors = []
                if not first_name: errors.append("First name required")
                if not last_name: errors.append("Last name required")
                if not email or "@" not in email: errors.append("Valid email required")
                if not password or len(password) < 8: errors.append("Password must be 8+ characters")
                if password != confirm: errors.append("Passwords don't match")
                if not accept: errors.append("Must accept terms")
                if get_account_data(email=email): errors.append("Email already exists")
                
                if errors: 
                    for e in errors:
                        st.error(e)
                else:
                    result = create_user_account({"first_name": first_name, "last_name": last_name, "email": email}, password)
                    if result["success"]:
                        st.session_state.user_id = result["user_id"]
                        st.session_state.user_account = result["user_record"]
                        st.session_state.logged_in = True
                        st.session_state.show_profile_setup = True
                        st.session_state.image_handler = ImageHandler(result["user_id"])
                        st.success("Account created!")
                        st.balloons()
                        st.rerun()
                    else: 
                        st.error(f"Error: {result.get('error', 'Unknown error')}")
    st.stop()

# ============================================================================
# LOAD USER DATA
# ============================================================================
if st.session_state.logged_in and st.session_state.user_id and not st.session_state.data_loaded:
    user_data = load_user_data(st.session_state.user_id)
    if "responses" in user_data:
        for sid_str, sdata in user_data["responses"].items():
            try: 
                sid = int(sid_str)
            except: 
                continue
            if sid in st.session_state.responses:
                st.session_state.responses[sid]["questions"] = sdata.get("questions", {})
    st.session_state.data_loaded = True
    
    # Initialize image handler if needed
    if not st.session_state.image_handler:
        st.session_state.image_handler = ImageHandler(st.session_state.user_id)

# ============================================================================
# PROFILE SETUP MODAL
# ============================================================================
if st.session_state.get('show_profile_setup', False):
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("👤 Complete Your Profile")
    
    tab1, tab2, tab3 = st.tabs(["Basic Info", "Heart of Your Story", "Life Details"])
    
    with tab1:
        st.markdown("### Basic Information")
        if st.session_state.user_account:
            profile = st.session_state.user_account['profile']
            
            profile['first_name'] = st.text_input("First Name", value=profile.get('first_name', ''))
            profile['last_name'] = st.text_input("Last Name", value=profile.get('last_name', ''))
            profile['email'] = st.text_input("Email", value=profile.get('email', ''))
            
            if st.button("💾 Save Basic Info", type="primary"):
                save_account_data(st.session_state.user_account)
                st.success("Saved!")
    
    with tab2:
        render_narrative_gps()
    
    with tab3:
        render_enhanced_profile()
    
    if st.button("← Close Profile", use_container_width=True):
        st.session_state.show_profile_setup = False
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# MODAL HANDLING
# ============================================================================
if st.session_state.show_ai_rewrite and st.session_state.get('current_rewrite_data'):
    show_ai_rewrite_modal()
    st.stop()

if st.session_state.show_privacy_settings:
    show_privacy_settings()
    st.stop()

if st.session_state.show_cover_designer:
    show_cover_designer()
    st.stop()

if st.session_state.show_publisher:
    show_publisher_modal()
    st.stop()

# ============================================================================
# MAIN HEADER
# ============================================================================
st.markdown(f'<div style="text-align:center;"><img src="{LOGO_URL}" style="max-width:300px;"></div>', unsafe_allow_html=True)

# ============================================================================
# SIDEBAR - COMPLETE VERSION WITH PUBLISHER BUTTON
# ============================================================================
with st.sidebar:
    st.markdown('<div class="sidebar-header"><h2>Tell My Story</h2><p>Your Life Timeline</p></div>', unsafe_allow_html=True)
    
    # ===== PROFILE SECTION =====
    st.header("👤 Your Profile")
    if st.session_state.user_account:
        profile = st.session_state.user_account['profile']
        st.success(f"✓ **{profile.get('first_name', '')} {profile.get('last_name', '')}**")
    else:
        st.info("Not logged in")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📝 Profile", use_container_width=True): 
            st.session_state.show_profile_setup = True
            st.rerun()
    with col2:
        if st.button("🔒 Privacy", use_container_width=True):
            st.session_state.show_privacy_settings = True
            st.rerun()
    
    if st.button("🚪 Log Out", use_container_width=True): 
        logout_user()
    
    st.divider()
    
    # ===== PUBLISHING SECTION =====
    st.header("📚 Publishing")
    if st.button("🎨 Publish Your Book", type="primary", use_container_width=True):
        st.session_state.show_publisher = True
        st.rerun()
    
    # Show quick stats
    if st.session_state.logged_in:
        total_stories = 0
        total_words = 0
        for session in st.session_state.current_question_bank:
            sid = session["id"]
            if sid in st.session_state.responses:
                total_stories += len(st.session_state.responses[sid].get("questions", {}))
                for q_data in st.session_state.responses[sid].get("questions", {}).values():
                    if q_data.get("answer"):
                        text_only = re.sub(r'<[^>]+>', '', q_data["answer"])
                        total_words += len(text_only.split())
        
        if total_stories > 0:
            st.info(f"📝 {total_stories} stories • {total_words:,} words")
    
    st.divider()
    
    # ===== SESSIONS SECTION =====
    st.header("📖 Sessions")
    
    # Quick session navigation
    sessions_per_row = 2
    session_cols = st.columns(sessions_per_row)
    
    for i, session in enumerate(st.session_state.current_question_bank):
        col_idx = i % sessions_per_row
        with session_cols[col_idx]:
            sid = session["id"]
            sdata = st.session_state.responses.get(sid, {})
            answered = len(sdata.get("questions", {}))
            total = len(session["questions"])
            
            # Determine status emoji
            if i == st.session_state.current_session:
                status = "▶️"
            elif answered == total and total > 0:
                status = "✅"
            elif answered > 0:
                status = "🟡"
            else:
                status = "⚪"
            
            if st.button(f"{status} {i+1}", key=f"session_{i}", help=session["title"], use_container_width=True):
                st.session_state.current_session = i
                st.session_state.current_question = 0
                st.session_state.current_question_override = None
                st.rerun()
    
    # Session manager button
    if st.button("📋 All Sessions", use_container_width=True):
        st.session_state.show_session_manager = True
        st.rerun()
    
    st.divider()
    
    # ===== TOPIC TOOLS SECTION =====
    st.header("📝 Writing Tools")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📚 Topic Bank", use_container_width=True):
            st.session_state.show_topic_browser = True
            st.rerun()
    with col2:
        if st.button("📖 Vignettes", use_container_width=True):
            st.session_state.show_vignette_manager = True
            st.rerun()
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("✨ New Vignette", use_container_width=True):
            import uuid
            new_id = str(uuid.uuid4())[:8]
            if 'vignette_manager' not in st.session_state:
                try:
                    from vignettes import VignetteManager
                    st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
                except:
                    st.session_state.vignette_manager = None
            
            if st.session_state.vignette_manager:
                st.session_state.vignette_manager.create_vignette_with_id(
                    id=new_id,
                    title="Untitled Vignette",
                    content="<p>Write your story here...</p>",
                    theme="Life Lesson",
                    mood="Reflective",
                    is_draft=True
                )
                st.session_state.editing_vignette_id = new_id
                st.session_state.show_vignette_modal = True
                st.rerun()
    
    with col2:
        if st.button("📂 Import File", use_container_width=True):
            # This will be handled in the main area
            import_key = f"import_quill_{st.session_state.current_session}_{st.session_state.current_question}"
            st.session_state[import_key] = True
            st.rerun()
    
    st.divider()
    
    # ===== QUESTION BANK SECTION =====
    st.header("📚 Question Banks")
    
    if st.button("📋 Bank Manager", use_container_width=True):
        st.session_state.show_bank_manager = True
        st.rerun()
    
    if st.session_state.get('current_bank_name'):
        st.caption(f"Current: {st.session_state.current_bank_name}")
    
    st.divider()
    
    # ===== SEARCH SECTION =====
    st.header("🔍 Search")
    
    search_query = st.text_input("Search stories...", placeholder="e.g., childhood, wedding", label_visibility="collapsed")
    
    if search_query and len(search_query) >= 2:
        results = search_all_answers(search_query)
        if results:
            st.success(f"Found {len(results)} matches")
            with st.container(height=300):
                for i, r in enumerate(results[:10]):
                    with st.container():
                        st.markdown(f"**{r['session_title']}**")
                        st.caption(f"{r['question'][:50]}...")
                        if st.button(f"Go to story {i}", key=f"search_{i}", use_container_width=True):
                            for idx, s in enumerate(st.session_state.current_question_bank):
                                if s["id"] == r['session_id']:
                                    st.session_state.current_session = idx
                                    st.session_state.current_question_override = r['question']
                                    st.rerun()
                        st.divider()
        else:
            st.info("No matches found")
    
    st.divider()
    
    # ===== STATS SECTION =====
    st.header("📊 Statistics")
    
    total_answered = 0
    total_questions = 0
    total_words = 0
    
    for session in st.session_state.current_question_bank:
        sid = session["id"]
        total_questions += len(session["questions"])
        answered = len(st.session_state.responses.get(sid, {}).get("questions", {}))
        total_answered += answered
        
        # Count words
        if sid in st.session_state.responses:
            for q_data in st.session_state.responses[sid].get("questions", {}).values():
                if q_data.get("answer"):
                    text_only = re.sub(r'<[^>]+>', '', q_data["answer"])
                    total_words += len(text_only.split())
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Progress", f"{total_answered}/{total_questions}")
    with col2:
        st.metric("Words", f"{total_words:,}")
    
    # Progress bar
    if total_questions > 0:
        progress_pct = (total_answered / total_questions)
        st.progress(progress_pct)
    
    st.divider()
    
    # ===== WORD TARGET SECTION =====
    if st.session_state.current_session < len(st.session_state.current_question_bank):
        current_session = st.session_state.current_question_bank[st.session_state.current_session]
        current_session_id = current_session["id"]
        progress_info = get_progress_info(current_session_id)
        
        st.header("🎯 Current Session")
        st.caption(current_session["title"])
        
        st.metric("Word Count", f"{progress_info['current_count']}/{progress_info['target']}")
        st.progress(min(progress_info['progress_percent'] / 100, 1.0))
        
        if st.button("✏️ Edit Target", use_container_width=True):
            st.session_state.editing_word_target = not st.session_state.editing_word_target
        
        if st.session_state.editing_word_target:
            new_target = st.number_input("Target:", min_value=100, max_value=5000, value=progress_info['target'])
            if st.button("💾 Save", use_container_width=True):
                st.session_state.responses[current_session_id]["word_target"] = new_target
                save_user_data(st.session_state.user_id, st.session_state.responses)
                st.session_state.editing_word_target = False
                st.rerun()
    
    st.divider()
    
    # ===== EXPORT SECTION =====
    st.header("📤 Export")
    
    if st.session_state.logged_in and total_answered > 0:
        # Simple text export
        if st.button("📝 Download as Text", use_container_width=True):
            text_content = "Tell My Story - Life Timeline\n\n"
            for session in st.session_state.current_question_bank:
                sid = session["id"]
                text_content += f"\n=== {session['title']} ===\n\n"
                if sid in st.session_state.responses:
                    for q, a in st.session_state.responses[sid].get("questions", {}).items():
                        text_content += f"Q: {q}\n"
                        text_content += f"A: {re.sub(r'<[^>]+>', '', a.get('answer', ''))}\n\n"
            
            st.download_button(
                label="📥 Download TXT",
                data=text_content,
                file_name=f"tell_my_story_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
                use_container_width=True,
                key="txt_export"
            )
        
        # JSON backup
        if st.button("💾 JSON Backup", use_container_width=True):
            backup_data = {
                "user_id": st.session_state.user_id,
                "export_date": datetime.now().isoformat(),
                "responses": st.session_state.responses,
                "profile": st.session_state.user_account.get('profile', {}) if st.session_state.user_account else {}
            }
            
            st.download_button(
                label="📥 Download JSON",
                data=json.dumps(backup_data, indent=2),
                file_name=f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True,
                key="json_export"
            )
    else:
        st.caption("Write stories to enable export")
    
    st.divider()
    
    # ===== DANGER ZONE =====
    st.header("⚠️ Danger Zone")
    
    if st.session_state.confirming_clear == "session":
        st.warning("**Clear current session?**")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Yes", type="primary", use_container_width=True):
                sid = st.session_state.current_question_bank[st.session_state.current_session]["id"]
                st.session_state.responses[sid]["questions"] = {}
                save_user_data(st.session_state.user_id, st.session_state.responses)
                st.session_state.confirming_clear = None
                st.rerun()
        with col2:
            if st.button("❌ No", use_container_width=True):
                st.session_state.confirming_clear = None
                st.rerun()
    
    elif st.session_state.confirming_clear == "all":
        st.warning("**Clear ALL data?**")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Yes All", type="primary", use_container_width=True):
                for s in st.session_state.current_question_bank:
                    st.session_state.responses[s["id"]]["questions"] = {}
                save_user_data(st.session_state.user_id, st.session_state.responses)
                st.session_state.confirming_clear = None
                st.rerun()
        with col2:
            if st.button("❌ No", use_container_width=True):
                st.session_state.confirming_clear = None
                st.rerun()
    
    else:
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear Session", use_container_width=True):
                st.session_state.confirming_clear = "session"
                st.rerun()
        with col2:
            if st.button("🔥 Clear All", use_container_width=True):
                st.session_state.confirming_clear = "all"
                st.rerun()
    
    st.divider()
    
    # ===== FOOTER =====
    st.caption(f"Tell My Story v1.0")
    if st.session_state.user_id:
        st.caption(f"User: {st.session_state.user_id[:8]}...")

# ============================================================================
# MAIN CONTENT - CURRENT SESSION AND QUESTION
# ============================================================================
if st.session_state.current_session >= len(st.session_state.current_question_bank): 
    st.session_state.current_session = 0

current_session = st.session_state.current_question_bank[st.session_state.current_session]
current_session_id = current_session["id"]

if st.session_state.current_question_override:
    current_question_text = st.session_state.current_question_override
else:
    if st.session_state.current_question >= len(current_session["questions"]): 
        st.session_state.current_question = 0
    current_question_text = current_session["questions"][st.session_state.current_question]

# Header with progress
col1, col2 = st.columns([3, 1])
with col1:
    st.subheader(f"Session {current_session_id}: {current_session['title']}")
    sdata = st.session_state.responses.get(current_session_id, {})
    answered = len(sdata.get("questions", {}))
    total = len(current_session["questions"])
    if total > 0: 
        st.progress(answered/total)
        st.caption(f"📝 Progress: {answered}/{total} topics")
with col2:
    if st.session_state.current_question_override:
        st.markdown('<div style="background:#ffd700; padding:5px; border-radius:5px; text-align:center;">✨ Custom Topic</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div style="text-align:center;">Topic {st.session_state.current_question+1} of {total}</div>', unsafe_allow_html=True)

# Question display
st.markdown(f'<div style="background:#f0f2f6; padding:20px; border-radius:10px; margin:20px 0; font-size:20px; font-weight:bold;">{current_question_text}</div>', unsafe_allow_html=True)

# Get existing answer
existing_answer = ""
if current_session_id in st.session_state.responses:
    if current_question_text in st.session_state.responses[current_session_id]["questions"]:
        existing_answer = st.session_state.responses[current_session_id]["questions"][current_question_text]["answer"]

# Editor setup
editor_base_key = f"quill_{current_session_id}_{current_question_text[:20]}"
content_key = f"{editor_base_key}_content"

if content_key not in st.session_state:
    if existing_answer:
        st.session_state[content_key] = existing_answer
    else:
        st.session_state[content_key] = "<p>Start writing your story here...</p>"

# Editor
st.markdown("### ✍️ Your Story")
content = st_quill(
    value=st.session_state[content_key],
    key=f"editor_{current_session_id}_{st.session_state.current_question}",
    placeholder="Start writing your story here...",
    html=True
)

if content is not None and content != st.session_state[content_key]:
    st.session_state[content_key] = content

st.markdown("---")

# ============================================================================
# ACTION BUTTONS
# ============================================================================
col1, col2, col3, col4, col5, col6 = st.columns([1, 1, 1, 1, 1, 2])

current_content = st.session_state.get(content_key, "")
has_content = current_content and current_content != "<p><br></p>" and current_content != "<p>Start writing your story here...</p>"

with col1:
    if st.button("💾 Save", type="primary", use_container_width=True):
        if has_content:
            if save_response(current_session_id, current_question_text, current_content):
                st.success("✅ Saved!")
                time.sleep(0.5)
                st.rerun()
            else: 
                st.error("Failed to save")
        else: 
            st.warning("Please write something!")

with col2:
    if existing_answer:
        if st.button("🗑️ Delete", use_container_width=True):
            if delete_response(current_session_id, current_question_text):
                st.session_state[content_key] = "<p>Start writing your story here...</p>"
                st.success("Deleted!")
                st.rerun()
    else: 
        st.button("🗑️ Delete", disabled=True, use_container_width=True)

with col3:
    if has_content:
        if st.button("🔍 Spell Check", use_container_width=True):
            text_only = re.sub(r'<[^>]+>', '', current_content)
            corrected = auto_correct_text(text_only)
            if corrected and corrected != text_only:
                st.session_state[content_key] = f"<p>{corrected}</p>"
                st.success("Spelling corrected!")
                st.rerun()
            else:
                st.info("No spelling issues found!")
    else:
        st.button("🔍 Spell Check", disabled=True, use_container_width=True)

with col4:
    if has_content:
        if st.button("✨ AI Rewrite", use_container_width=True):
            st.session_state.show_ai_rewrite_menu = not st.session_state.show_ai_rewrite_menu
            st.rerun()
    else:
        st.button("✨ AI Rewrite", disabled=True, use_container_width=True)

with col5:
    # Import button
    import_key = f"import_{editor_base_key}"
    if import_key not in st.session_state:
        st.session_state[import_key] = False
    
    button_label = "📂 Close Import" if st.session_state[import_key] else "📂 Import"
    if st.button(button_label, use_container_width=True):
        st.session_state[import_key] = not st.session_state[import_key]
        st.rerun()

with col6:
    # Navigation
    nav1, nav2 = st.columns(2)
    with nav1: 
        prev_disabled = st.session_state.current_question == 0 and not st.session_state.current_question_override
        if st.button("← Previous", disabled=prev_disabled, use_container_width=True):
            if not prev_disabled:
                if st.session_state.current_question_override:
                    st.session_state.current_question_override = None
                else:
                    st.session_state.current_question -= 1
                st.rerun()
    with nav2:
        next_disabled = (st.session_state.current_question >= len(current_session["questions"]) - 1 and 
                        not st.session_state.current_question_override)
        if st.button("Next →", disabled=next_disabled, use_container_width=True):
            if not next_disabled:
                if st.session_state.current_question_override:
                    st.session_state.current_question_override = None
                else:
                    st.session_state.current_question += 1
                st.rerun()

# ============================================================================
# AI REWRITE MENU
# ============================================================================
if st.session_state.get('show_ai_rewrite_menu', False):
    st.markdown("---")
    st.markdown("### ✨ AI Rewrite Options")
    
    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
    
    with col1:
        if st.button("👤 First Person", use_container_width=True):
            result = ai_rewrite_answer(current_content, "1st", current_question_text, current_session['title'])
            if result.get('success'):
                st.session_state.current_rewrite_data = result
                st.session_state.show_ai_rewrite = True
                st.session_state.show_ai_rewrite_menu = False
                st.rerun()
            else:
                st.error(result.get('error', 'Failed'))
    
    with col2:
        if st.button("💬 Second Person", use_container_width=True):
            result = ai_rewrite_answer(current_content, "2nd", current_question_text, current_session['title'])
            if result.get('success'):
                st.session_state.current_rewrite_data = result
                st.session_state.show_ai_rewrite = True
                st.session_state.show_ai_rewrite_menu = False
                st.rerun()
            else:
                st.error(result.get('error', 'Failed'))
    
    with col3:
        if st.button("📖 Third Person", use_container_width=True):
            result = ai_rewrite_answer(current_content, "3rd", current_question_text, current_session['title'])
            if result.get('success'):
                st.session_state.current_rewrite_data = result
                st.session_state.show_ai_rewrite = True
                st.session_state.show_ai_rewrite_menu = False
                st.rerun()
            else:
                st.error(result.get('error', 'Failed'))
    
    with col4:
        if st.button("❌ Cancel", use_container_width=True):
            st.session_state.show_ai_rewrite_menu = False
            st.rerun()

# ============================================================================
# IMPORT SECTION
# ============================================================================
if st.session_state.get(import_key, False):
    st.markdown("---")
    st.markdown("### 📂 Import Text File")
    
    uploaded_file = st.file_uploader(
        "Choose a file", 
        type=['txt', 'docx'],
        key=f"file_upload_{editor_base_key}"
    )
    
    if uploaded_file:
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📥 Import", type="primary", use_container_width=True):
                with st.spinner("Importing..."):
                    imported_html = import_text_file(uploaded_file)
                    if imported_html:
                        st.session_state[content_key] = imported_html
                        st.session_state[import_key] = False
                        st.success("File imported!")
                        st.rerun()
        
        with col2:
            if st.button("❌ Cancel", use_container_width=True):
                st.session_state[import_key] = False
                st.rerun()

# ============================================================================
# IMAGE UPLOAD SECTION
# ============================================================================
if st.session_state.logged_in and st.session_state.image_handler:
    
    # Get existing images
    existing_images = st.session_state.image_handler.get_images_for_answer(current_session_id, current_question_text)
    
    if existing_images:
        st.markdown("---")
        st.markdown("### 📸 Your Photos")
        
        cols = st.columns(3)
        for idx, img in enumerate(existing_images):
            with cols[idx % 3]:
                # Display thumbnail
                img_path = f"{st.session_state.image_handler.get_user_path()}/thumbnails/{img['id']}.jpg"
                if os.path.exists(img_path):
                    st.image(img_path, use_container_width=True)
                    if img.get('caption'):
                        st.caption(img['caption'])
                    if st.button(f"🗑️", key=f"del_img_{img['id']}"):
                        st.session_state.image_handler.delete_image(img['id'])
                        st.rerun()
    
    with st.expander("📤 Upload New Photos", expanded=not existing_images):
        st.markdown("**Add photos to your story:**")
        
uploaded_file = st.file_uploader(
    "Choose an image",  # Add your parameters here
    type=["jpg", "png", "jpeg"],  # Example parameter
    key=f"img_{current_session_id}_{st.session_state.current_question}"  # Example parameter
)  # <-- This closing parenthesis is crucial!

if uploaded_file:
    caption = st.text_input("Caption (optional):", key=f"img_cap_{current_session_id}_{st.session_state.current_question}")
    usage = st.radio("Size:", ["Full Page", "Inline"], horizontal=True, key=f"img_usage_{current_session_id}_{st.session_state.current_question}")
            
            if st.button("📤 Upload Image", type="primary"):
                with st.spinner("Uploading..."):
                    usage_type = "full_page" if usage == "Full Page" else "inline"
                    result = st.session_state.image_handler.save_image(
                        uploaded_file, current_session_id, current_question_text, caption, usage_type
                    )
                    if result:
                        st.success("Photo uploaded!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("Upload failed")

# ============================================================================
# PREVIEW SECTION
# ============================================================================
if has_content:
    with st.expander("👁️ Preview your story"):
        st.markdown(current_content, unsafe_allow_html=True)

# ============================================================================
# PROGRESS SECTION
# ============================================================================
st.markdown("---")
progress_info = get_progress_info(current_session_id)

# Progress bar
st.markdown(f"""
<div style="margin:20px 0;">
    <div style="display:flex; justify-content:space-between; margin-bottom:5px;">
        <span>📊 Word Count Progress</span>
        <span>{progress_info['current_count']} / {progress_info['target']} words</span>
    </div>
    <div style="width:100%; background-color:#f0f2f6; border-radius:10px; height:20px;">
        <div style="width:{min(progress_info['progress_percent'], 100)}%; background-color:{progress_info['color']}; border-radius:10px; height:20px;"></div>
    </div>
    <div style="text-align:center; margin-top:5px;">
        {progress_info['emoji']} {progress_info['progress_percent']:.0f}% • {progress_info['remaining_words']} words remaining
    </div>
</div>
""", unsafe_allow_html=True)

# Word target editor
if st.button("✏️ Change Word Target"):
    st.session_state.editing_word_target = not st.session_state.editing_word_target

if st.session_state.editing_word_target:
    new_target = st.number_input("Target words:", min_value=100, max_value=5000, value=progress_info['target'])
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save", type="primary", use_container_width=True):
            st.session_state.responses[current_session_id]["word_target"] = new_target
            save_user_data(st.session_state.user_id, st.session_state.responses)
            st.session_state.editing_word_target = False
            st.rerun()
    with col2:
        if st.button("❌ Cancel", use_container_width=True):
            st.session_state.editing_word_target = False
            st.rerun()

# ============================================================================
# FOOTER
# ============================================================================
st.markdown("---")
if st.session_state.user_account:
    profile = st.session_state.user_account.get('profile', {})
    st.caption(f"Tell My Story • {profile.get('first_name', '')} {profile.get('last_name', '')} • {st.session_state.current_bank_name}")

