# vignettes.py - COMPLETE WORKING VERSION WITH FILE IMPORT AND CLEAR CONTENT FIX
import streamlit as st
import json
from datetime import datetime
import os
import uuid
import re
import base64
import hashlib
import time
import openai

from streamlit_quill import st_quill

class VignetteManager:
    def __init__(self, user_id):
        self.user_id = user_id
        self.file = f"user_vignettes/{user_id}_vignettes.json"
        os.makedirs("user_vignettes", exist_ok=True)
        os.makedirs(f"user_vignettes/{user_id}_images", exist_ok=True)
        self.standard_themes = [
            "Life Lesson", "Achievement", "Work Experience", "Loss of Life",
            "Illness", "New Child", "Marriage", "Travel", "Relationship",
            "Interests", "Education", "Childhood Memory", "Family Story",
            "Career Moment", "Personal Growth"
        ]
        self._load()
    
    def _load(self):
        try:
            if os.path.exists(self.file):
                with open(self.file, 'r') as f:
                    self.vignettes = json.load(f)
            else:
                self.vignettes = []
        except:
            self.vignettes = []
    
    def _save(self):
        with open(self.file, 'w') as f:
            json.dump(self.vignettes, f, indent=2)
    
    def save_vignette_image(self, uploaded_file, vignette_id):
        try:
            file_ext = uploaded_file.name.split('.')[-1].lower()
            image_id = hashlib.md5(f"{vignette_id}{uploaded_file.name}{datetime.now()}".encode()).hexdigest()[:12]
            filename = f"{image_id}.{file_ext}"
            filepath = f"user_vignettes/{self.user_id}_images/{filename}"
            
            with open(filepath, 'wb') as f:
                f.write(uploaded_file.getbuffer())
            
            img_bytes = uploaded_file.getvalue()
            img_base64 = base64.b64encode(img_bytes).decode()
            
            return {
                "id": image_id,
                "filename": filename,
                "base64": img_base64,
                "path": filepath,
                "caption": ""
            }
        except Exception as e:
            st.error(f"Error saving image: {e}")
            return None
    
    def create_vignette(self, title, content, theme, mood="Reflective", is_draft=False, images=None):
        v = {
            "id": str(uuid.uuid4())[:8],
            "title": title,
            "content": content,
            "theme": theme,
            "mood": mood,
            "word_count": len(re.sub(r'<[^>]+>', '', content).split()),
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "is_draft": is_draft,
            "is_published": not is_draft,
            "images": images or []
        }
        self.vignettes.append(v)
        self._save()
        return v
    
    def update_vignette(self, id, title, content, theme, mood=None, images=None):
        for v in self.vignettes:
            if v["id"] == id:
                v.update({
                    "title": title, 
                    "content": content, 
                    "theme": theme, 
                    "mood": mood or v.get("mood", "Reflective"),
                    "word_count": len(re.sub(r'<[^>]+>', '', content).split()), 
                    "updated_at": datetime.now().isoformat(),
                    "images": images or v.get("images", [])
                })
                self._save()
                return True
        return False
    
    def delete_vignette(self, id):
        self.vignettes = [v for v in self.vignettes if v["id"] != id]
        self._save()
        return True
    
    def get_vignette_by_id(self, id):
        for v in self.vignettes:
            if v["id"] == id:
                return v
        return None
    
    def check_spelling(self, text):
        """Check spelling and grammar using OpenAI"""
        if not text: 
            return text
        try:
            client = openai.OpenAI(api_key=st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY")))
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Fix spelling and grammar. Return only corrected text."},
                    {"role": "user", "content": text}
                ],
                max_tokens=len(text) + 100, 
                temperature=0.1
            )
            return resp.choices[0].message.content
        except Exception as e:
            st.error(f"Spell check failed: {e}")
            return text
    
    def ai_rewrite_vignette(self, original_text, person_option, vignette_title):
        """Rewrite the vignette in 1st, 2nd, or 3rd person using profile context"""
        try:
            # Get OpenAI client
            client = openai.OpenAI(api_key=st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY")))
            
            # Get profile context from user account in session state
            gps_context = ""
            enhanced_context = ""
            
            if st.session_state.get('user_account'):
                # Narrative GPS context
                if 'narrative_gps' in st.session_state.user_account:
                    gps = st.session_state.user_account['narrative_gps']
                    if gps:
                        gps_context = "\n\n=== BOOK PROJECT CONTEXT ===\n"
                        if gps.get('book_title'): gps_context += f"- Book Title: {gps['book_title']}\n"
                        if gps.get('genre'): 
                            genre = gps['genre']
                            if genre == "Other" and gps.get('genre_other'):
                                genre = gps['genre_other']
                            gps_context += f"- Genre: {genre}\n"
                        if gps.get('purposes'): 
                            gps_context += f"- Purpose: {', '.join(gps['purposes'])}\n"
                        if gps.get('narrative_voices'): 
                            gps_context += f"- Narrative Voice: {', '.join(gps['narrative_voices'])}\n"
                        if gps.get('emotional_tone'): 
                            gps_context += f"- Emotional Tone: {gps['emotional_tone']}\n"
                        if gps.get('reader_takeaway'): 
                            gps_context += f"- Reader Takeaway: {gps['reader_takeaway']}\n"
                
                # Enhanced profile context
                if 'enhanced_profile' in st.session_state.user_account:
                    ep = st.session_state.user_account['enhanced_profile']
                    if ep:
                        enhanced_context = "\n\n=== BIOGRAPHER CONTEXT ===\n"
                        if ep.get('birth_place'): enhanced_context += f"• Born: {ep['birth_place']}\n"
                        if ep.get('parents'): enhanced_context += f"• Parents: {ep['parents'][:150]}...\n"
                        if ep.get('childhood_home'): enhanced_context += f"• Childhood: {ep['childhood_home'][:150]}...\n"
                        if ep.get('career_path'): enhanced_context += f"• Career: {ep['career_path'][:150]}...\n"
                        if ep.get('life_lessons'): enhanced_context += f"• Life Philosophy: {ep['life_lessons'][:200]}...\n"
                        if ep.get('legacy'): enhanced_context += f"• Legacy Hope: {ep['legacy'][:200]}...\n"
            
            # Clean the text (remove HTML tags)
            clean_text = re.sub(r'<[^>]+>', '', original_text)
            
            if len(clean_text.split()) < 5:
                return {"error": "Text too short to rewrite (minimum 5 words)"}
            
            # Person-specific instructions
            person_instructions = {
                "1st": {
                    "name": "First Person",
                    "instruction": "Rewrite this in FIRST PERSON ('I', 'me', 'my', 'we', 'our'). Keep the authentic voice of the author telling their own story.",
                    "example": "I remember the day clearly. The sun was setting and I felt...",
                    "emoji": "👤"
                },
                "2nd": {
                    "name": "Second Person",
                    "instruction": "Rewrite this in SECOND PERSON ('you', 'your') as if speaking directly to the reader. Make it feel like advice, a letter, or a conversation with the reader.",
                    "example": "You remember that day clearly. The sun was setting and you felt...",
                    "emoji": "💬"
                },
                "3rd": {
                    "name": "Third Person",
                    "instruction": "Rewrite this in THIRD PERSON ('he', 'she', 'they', 'the author', the person's name). Write as if telling someone else's story to readers.",
                    "example": "They remember the day clearly. The sun was setting and they felt...",
                    "emoji": "📖"
                }
            }
            
            # Get author's name for 3rd person
            author_name = ""
            if st.session_state.get('user_account'):
                profile = st.session_state.user_account.get('profile', {})
                first = profile.get('first_name', '')
                last = profile.get('last_name', '')
                if first and last:
                    author_name = f"{first} {last}"
                elif first:
                    author_name = first
            
            system_prompt = f"""You are an expert writing assistant and ghostwriter. Your task is to rewrite this vignette in {person_instructions[person_option]['name']}.

{person_instructions[person_option]['instruction']}

EXAMPLE STYLE:
{person_instructions[person_option]['example']}

IMPORTANT GUIDELINES:
1. Use the profile context below to understand WHO the author is
2. Preserve all key facts, emotions, and details from the original
3. Maintain the author's unique voice and personality
4. Fix any grammar issues naturally
5. Make it flow better while keeping it authentic
6. DO NOT add fictional events or details not in the original
7. If using third person and you know the author's name, use it naturally
8. Return ONLY the rewritten text, no explanations, no prefixes

PROFILE CONTEXT (Use this to understand the author's voice):
{gps_context}
{enhanced_context}

VIGNETTE TITLE: {vignette_title}
AUTHOR NAME: {author_name if author_name else 'Unknown'}

ORIGINAL VIGNETTE (to rewrite):
{clean_text}

REWRITTEN VERSION ({person_instructions[person_option]['name']}):"""

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": "Please rewrite this vignette in the specified voice."}
                ],
                max_tokens=len(clean_text.split()) * 3,
                temperature=0.7
            )
            
            rewritten = response.choices[0].message.content.strip()
            
            # Clean up any markdown or quotes the AI might add
            rewritten = re.sub(r'^["\']|["\']$', '', rewritten)
            rewritten = re.sub(r'^Here\'s the rewritten version:?\s*', '', rewritten, flags=re.IGNORECASE)
            
            return {
                "success": True,
                "original": clean_text,
                "rewritten": rewritten,
                "person": person_instructions[person_option]["name"],
                "emoji": person_instructions[person_option]["emoji"]
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    # ============================================================================
    # FILE IMPORT FUNCTION FOR VIGNETTES
    # ============================================================================
    def import_text_file(self, uploaded_file):
        """Import text from common document formats into vignette"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            file_content = ""
            
            # Get file size in MB
            file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
            
            # Show what's being imported
            st.info(f"📄 **Importing:** {uploaded_file.name}")
            st.caption(f"Size: {file_size_mb:.1f}MB • Format: .{file_extension}")
            
            # ============================================
            # SUPPORTED FORMATS
            # ============================================
            
            # 1. PLAIN TEXT - Universal format
            if file_extension == 'txt':
                file_content = uploaded_file.read().decode('utf-8', errors='ignore')
                st.success(f"✅ Plain text file loaded ({len(file_content.split())} words)")
            
            # 2. MICROSOFT WORD - Critical for business docs, interviews, transcripts
            elif file_extension == 'docx':
                try:
                    import io
                    from docx import Document
                    docx_bytes = io.BytesIO(uploaded_file.getvalue())
                    doc = Document(docx_bytes)
                    
                    # Extract paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    
                    # Extract from tables (useful for interview transcripts)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' '.join([cell.text for cell in row.cells if cell.text.strip()])
                            if row_text:
                                paragraphs.append(row_text)
                    
                    file_content = '\n\n'.join(paragraphs)
                    st.success(f"✅ Word document loaded ({len(paragraphs)} paragraphs, {len(file_content.split())} words)")
                except ImportError:
                    st.error("📦 Missing dependency: Please install python-docx")
                    st.code("pip install python-docx")
                    return None
                except Exception as e:
                    st.error(f"Error reading Word document: {str(e)}")
                    return None
            
            # 3. RICH TEXT FORMAT - Sometimes used by transcription services
            elif file_extension == 'rtf':
                try:
                    from striprtf.striprtf import rtf_to_text
                    rtf_content = uploaded_file.read().decode('utf-8', errors='ignore')
                    file_content = rtf_to_text(rtf_content)
                    st.success(f"✅ RTF document loaded ({len(file_content.split())} words)")
                except ImportError:
                    st.warning("📦 Optional: For RTF support, install striprtf")
                    st.code("pip install striprtf")
                    return None
                except Exception as e:
                    st.error(f"Error reading RTF: {str(e)}")
                    return None
            
            # 4. SUBTITLE FILES - Common from automated transcription
            elif file_extension in ['vtt', 'srt']:
                file_content = uploaded_file.read().decode('utf-8', errors='ignore')
                # Clean up subtitle formatting
                lines = file_content.split('\n')
                clean_lines = []
                for line in lines:
                    # Remove timestamps and numbers
                    if '-->' not in line and not line.strip().isdigit() and line.strip():
                        clean_lines.append(line.strip())
                file_content = ' '.join(clean_lines)
                st.success(f"✅ Subtitle file loaded ({len(clean_lines)} lines, {len(file_content.split())} words)")
            
            # 5. JSON - Some transcription APIs export this format
            elif file_extension == 'json':
                try:
                    import json
                    data = json.loads(uploaded_file.read().decode('utf-8'))
                    
                    # Try common JSON structures from transcription services
                    if isinstance(data, dict):
                        # OpenAI Whisper / Google Speech-to-Text
                        if 'text' in data:
                            file_content = data['text']
                            st.success(f"✅ JSON transcript loaded ({len(file_content.split())} words)")
                        elif 'results' in data:
                            # Google Speech-to-Text format
                            texts = []
                            for r in data['results']:
                                if 'alternatives' in r and r['alternatives']:
                                    texts.append(r['alternatives'][0].get('transcript', ''))
                            file_content = ' '.join(texts)
                            st.success(f"✅ Google Speech-to-Text JSON loaded ({len(file_content.split())} words)")
                        elif 'transcript' in data:
                            # Otter.ai format
                            if isinstance(data['transcript'], list):
                                texts = [t.get('text', '') for t in data['transcript'] if t.get('text')]
                                file_content = ' '.join(texts)
                            else:
                                file_content = data['transcript']
                            st.success(f"✅ Otter.ai JSON loaded ({len(file_content.split())} words)")
                        else:
                            # Unknown structure - show preview
                            preview = str(data)[:200] + "..." if len(str(data)) > 200 else str(data)
                            st.warning(f"⚠️ Unknown JSON structure. Preview: {preview}")
                            file_content = str(data)
                    else:
                        file_content = str(data)
                except Exception as e:
                    st.error(f"Error parsing JSON: {str(e)}")
                    return None
            
            # 6. MARKDOWN - For notes and documentation
            elif file_extension == 'md':
                file_content = uploaded_file.read().decode('utf-8', errors='ignore')
                # Basic markdown cleanup (remove formatting symbols)
                file_content = re.sub(r'#{1,6}\s*', '', file_content)  # Remove headers
                file_content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', file_content)  # Remove links
                file_content = re.sub(r'[*_]{1,2}([^*_]+)[*_]{1,2}', r'\1', file_content)  # Remove emphasis
                st.success(f"✅ Markdown file loaded ({len(file_content.split())} words)")
            
            # 7. Unsupported format
            else:
                st.error(f"❌ Unsupported file format: .{file_extension}")
                st.info("Supported formats: .txt, .docx, .rtf, .vtt, .srt, .json, .md")
                return None
            
            # Validate we got content
            if not file_content or not file_content.strip():
                st.warning("The file appears to be empty or couldn't be read")
                return None
            
            # Check file size warning
            if file_size_mb > 10:
                st.warning(f"⚠️ Large file ({file_size_mb:.1f}MB) - processing may be slow")
            
            # Clean up the text
            # Remove excessive whitespace
            file_content = re.sub(r'\s+', ' ', file_content)
            
            # Split into paragraphs (by sentences for speech-to-text which often lacks paragraph breaks)
            sentences = re.split(r'[.!?]+', file_content)
            paragraphs = []
            current_para = []
            
            for sentence in sentences:
                if sentence.strip():
                    current_para.append(sentence.strip() + '.')
                    # Create a new paragraph every 3-4 sentences
                    if len(current_para) >= 4:
                        paragraphs.append(' '.join(current_para))
                        current_para = []
            
            # Add any remaining sentences
            if current_para:
                paragraphs.append(' '.join(current_para))
            
            # If no paragraphs created (very short text), use the whole thing
            if not paragraphs:
                paragraphs = [file_content]
            
            # Wrap each paragraph in <p> tags
            html_content = ''
            for para in paragraphs:
                if para.strip():
                    # Escape any HTML characters
                    para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    html_content += f'<p>{para.strip()}</p>'
            
            return html_content
            
        except Exception as e:
            st.error(f"Error importing file: {str(e)}")
            return None
    
    def display_vignette_creator(self, on_publish=None, edit_vignette=None):
        # Create STABLE keys for this vignette
        if edit_vignette:
            vignette_id = edit_vignette['id']
            base_key = f"vignette_{vignette_id}"
            is_new = False
        else:
            # For NEW vignette, use a timestamp to ensure unique keys
            import time
            vignette_id = f"new_{int(time.time())}"
            base_key = f"vignette_{vignette_id}"
            is_new = True
        
        # Editor key and content key
        editor_key = f"quill_vignette_{vignette_id}"
        content_key = f"{editor_key}_content"
        
        # Add a version counter for this editor
        version_key = f"{editor_key}_version"
        if version_key not in st.session_state:
            st.session_state[version_key] = 0
        
        # IMPORTANT: For NEW vignettes, clear any previous content
        if is_new:
            # Check if this is a brand new vignette (no previous state)
            if f"{base_key}_initialized" not in st.session_state:
                # Clear any existing content for this key
                if content_key in st.session_state:
                    del st.session_state[content_key]
                # Set initialized flag
                st.session_state[f"{base_key}_initialized"] = True
        
        # Title input
        title = st.text_input(
            "Title", 
            value=edit_vignette.get("title", "") if edit_vignette else "",
            placeholder="Give your vignette a meaningful title",
            key=f"{base_key}_title"
        )
        
        # Theme and mood in columns
        col1, col2 = st.columns(2)
        with col1:
            theme_options = self.standard_themes + ["Custom"]
            if edit_vignette and edit_vignette.get("theme"):
                current_theme = edit_vignette["theme"]
                if current_theme in self.standard_themes:
                    theme_index = self.standard_themes.index(current_theme)
                    theme = st.selectbox("Theme", theme_options, index=theme_index, key=f"{base_key}_theme")
                else:
                    theme = st.selectbox("Theme", theme_options, index=len(theme_options)-1, key=f"{base_key}_theme")
                    if theme == "Custom":
                        theme = st.text_input("Custom Theme", value=current_theme, key=f"{base_key}_custom_theme")
            else:
                theme = st.selectbox("Theme", theme_options, key=f"{base_key}_theme")
                if theme == "Custom":
                    theme = st.text_input("Custom Theme", key=f"{base_key}_custom_theme")
        
        with col2:
            mood_options = ["Reflective", "Joyful", "Bittersweet", "Humorous", "Serious", "Inspiring", "Nostalgic"]
            if edit_vignette:
                current_mood = edit_vignette.get("mood", "Reflective")
                mood_index = mood_options.index(current_mood) if current_mood in mood_options else 0
                mood = st.selectbox("Mood/Tone", mood_options, index=mood_index, key=f"{base_key}_mood")
            else:
                mood = st.selectbox("Mood/Tone", mood_options, key=f"{base_key}_mood")
        
        # Initialize content in session state
        if edit_vignette and edit_vignette.get("content"):
            default_content = edit_vignette["content"]
        else:
            default_content = "<p>Write your story here...</p>"
        
        if content_key not in st.session_state:
            st.session_state[content_key] = default_content
        
        st.markdown("### 📝 Your Story")
        st.markdown("""
        <div class="image-drop-info">
            📸 <strong>Drag & drop images</strong> directly into the editor.
        </div>
        """, unsafe_allow_html=True)
        
        # Editor component key with version
        editor_component_key = f"quill_editor_{vignette_id}_v{st.session_state[version_key]}"
        
        # Display Quill editor
        content = st_quill(
            value=st.session_state[content_key],
            key=editor_component_key,
            placeholder="Write your story here...",
            html=True
        )
        
        # Update session state when content changes
        if content is not None and content != st.session_state[content_key]:
            st.session_state[content_key] = content
        
        st.markdown("---")
        
        # ============================================================================
        # BUTTONS ROW - WITH IMPORT BUTTON
        # ============================================================================
        col1, col2, col3, col4, col5, col6, col7 = st.columns([1, 1, 1, 1, 1, 1, 2])
        
        # Spellcheck state management
        spellcheck_base = f"spell_{editor_key}"
        spell_result_key = f"{spellcheck_base}_result"
        current_content = st.session_state.get(content_key, "")
        has_content = current_content and current_content != "<p><br></p>" and current_content != "<p>Write your story here...</p>"
        showing_results = spell_result_key in st.session_state and st.session_state[spell_result_key].get("show", False)
        
        # Import state management
        import_key = f"import_{editor_key}"
        show_import = st.session_state.get(import_key, False)
        
        with col1:
            if st.button("💾 Save Draft", key=f"{base_key}_save_draft", type="primary", use_container_width=True):
                current_content = st.session_state[content_key]
                if not current_content or current_content == "<p><br></p>" or current_content == "<p></p>":
                    st.error("Please write some content")
                else:
                    final_title = title.strip() or "Untitled"
                    
                    if edit_vignette:
                        self.update_vignette(edit_vignette["id"], final_title, current_content, theme, mood)
                        st.success("✅ Draft saved!")
                        st.session_state.edit_success = True
                    else:
                        self.create_vignette(final_title, current_content, theme, mood, is_draft=True)
                        st.success("✅ Draft saved!")
                        st.session_state.draft_success = True
                    
                    if spell_result_key in st.session_state:
                        del st.session_state[spell_result_key]
                    
                    time.sleep(1)
                    st.session_state.show_vignette_modal = False
                    st.session_state.show_vignette_manager = True
                    st.rerun()
        
        with col2:
            if st.button("📢 Publish", key=f"{base_key}_publish", use_container_width=True, type="primary"):
                current_content = st.session_state[content_key]
                if not current_content or current_content == "<p><br></p>" or current_content == "<p></p>":
                    st.error("Please write some content")
                else:
                    final_title = title.strip() or "Untitled"
                    
                    if edit_vignette:
                        edit_vignette["is_draft"] = False
                        edit_vignette["published_at"] = datetime.now().isoformat()
                        self.update_vignette(edit_vignette["id"], final_title, current_content, theme, mood)
                        st.success("🎉 Published successfully!")
                        vignette_data = edit_vignette
                    else:
                        v = self.create_vignette(final_title, current_content, theme, mood, is_draft=False)
                        v["published_at"] = datetime.now().isoformat()
                        self.update_vignette(v["id"], final_title, current_content, theme, mood)
                        st.success("🎉 Published successfully!")
                        vignette_data = v
                    
                    if on_publish:
                        on_publish(vignette_data)
                    
                    if spell_result_key in st.session_state:
                        del st.session_state[spell_result_key]
                    
                    time.sleep(1)
                    st.session_state.show_vignette_modal = False
                    st.session_state.show_vignette_manager = True
                    st.rerun()
        
        with col3:
            if has_content and not showing_results:
                if st.button("🔍 Spell Check", key=f"{base_key}_spell", use_container_width=True):
                    with st.spinner("Checking spelling and grammar..."):
                        text_only = re.sub(r'<[^>]+>', '', current_content)
                        if len(text_only.split()) >= 3:
                            corrected = self.check_spelling(text_only)
                            if corrected and corrected != text_only:
                                st.session_state[spell_result_key] = {
                                    "original": text_only,
                                    "corrected": corrected,
                                    "show": True
                                }
                            else:
                                st.session_state[spell_result_key] = {
                                    "message": "✅ No spelling or grammar issues found!",
                                    "show": True
                                }
                            st.rerun()
                        else:
                            st.warning("Text too short for spell check (minimum 3 words)")
            else:
                st.button("🔍 Spell Check", key=f"{base_key}_spell_disabled", disabled=True, use_container_width=True)
        
        with col4:
            if has_content:
                if st.button("✨ AI Rewrite", key=f"{base_key}_ai_rewrite", use_container_width=True):
                    st.session_state[f"{base_key}_show_ai_menu"] = True
                    st.rerun()
            else:
                st.button("✨ AI Rewrite", key=f"{base_key}_ai_disabled", disabled=True, use_container_width=True)
        
        with col5:
            # IMPORT BUTTON
            if st.button("📂 Import File", key=f"{base_key}_import", use_container_width=True):
                st.session_state[import_key] = not st.session_state.get(import_key, False)
                st.rerun()
        
        with col6:
            if st.session_state.get(f"{base_key}_show_ai_menu", False):
                person_option = st.selectbox(
                    "Voice:",
                    options=["1st", "2nd", "3rd"],
                    format_func=lambda x: {"1st": "👤 First Person", "2nd": "💬 Second Person", "3rd": "📖 Third Person"}[x],
                    key=f"{base_key}_ai_person",
                    label_visibility="collapsed"
                )
                
                if st.button("Go", key=f"{base_key}_ai_go", type="primary", use_container_width=True):
                    with st.spinner(f"Rewriting in {person_option} person..."):
                        result = self.ai_rewrite_vignette(
                            current_content, 
                            person_option, 
                            title or "Untitled Vignette"
                        )
                        
                        if result.get('success'):
                            st.session_state[f"{base_key}_ai_result"] = result
                            st.session_state[f"{base_key}_show_ai_menu"] = False
                            st.rerun()
                        else:
                            st.error(result.get('error', 'Failed to rewrite'))
            else:
                st.markdown("")
        
        with col7:
            nav1, nav2 = st.columns(2)
            with nav1:
                if st.button("👁️ Preview", key=f"{base_key}_preview", use_container_width=True):
                    st.session_state[f"{base_key}_show_preview"] = True
                    st.rerun()
            with nav2:
                if st.button("❌ Cancel", key=f"{base_key}_cancel", use_container_width=True):
                    # Clear all session state for this vignette
                    keys_to_clear = [content_key, version_key, spell_result_key, 
                                    f"{base_key}_ai_result", f"{base_key}_show_ai_menu", 
                                    f"{base_key}_show_preview", import_key,
                                    f"{import_key}_pending", f"{import_key}_show_options"]
                    for key in keys_to_clear:
                        if key in st.session_state:
                            try:
                                del st.session_state[key]
                            except:
                                pass
                    st.session_state.show_vignette_modal = False
                    st.session_state.editing_vignette_id = None
                    st.rerun()
        
        # Display import section if toggled
        if show_import:
            st.markdown("---")
            st.markdown("### 📂 Import Text File")
            
            # Show supported formats table
            with st.expander("📋 Supported File Formats", expanded=True):
                st.markdown("""
                | Format | Description | Typical Use |
                |--------|-------------|-------------|
                | **.txt** | Plain text | Universal - works everywhere |
                | **.docx** | Microsoft Word | Business docs, interview transcripts |
                | **.rtf** | Rich Text Format | Older word processors |
                | **.vtt/.srt** | Subtitle files | Automated transcription output |
                | **.json** | JSON data | Speech-to-text APIs (Whisper, Google) |
                | **.md** | Markdown | Notes, documentation |
                
                **Maximum file size:** 50MB
                """)
            
            uploaded_file = st.file_uploader(
                "Choose a file to import",
                type=['txt', 'docx', 'rtf', 'vtt', 'srt', 'json', 'md'],
                key=f"{base_key}_file_uploader",
                help="Select a file from your computer to import into this vignette"
            )
            
            if uploaded_file:
                col_imp1, col_imp2, col_imp3 = st.columns([1, 1, 2])
                with col_imp1:
                    if st.button("📥 Import", key=f"{base_key}_do_import", type="primary", use_container_width=True):
                        with st.spinner("Importing file..."):
                            imported_html = self.import_text_file(uploaded_file)
                            if imported_html:
                                # Check if there's existing content
                                current = st.session_state.get(content_key, "")
                                if current and current != "<p>Write your story here...</p>" and current != "<p><br></p>":
                                    # Ask user what to do
                                    st.session_state[f"{import_key}_pending"] = imported_html
                                    st.session_state[f"{import_key}_show_options"] = True
                                    st.rerun()
                                else:
                                    # No existing content, just replace
                                    st.session_state[content_key] = imported_html
                                    st.session_state[version_key] += 1
                                    st.session_state[import_key] = False
                                    st.success("✅ File imported successfully!")
                                    st.rerun()
                
                with col_imp2:
                    if st.button("❌ Cancel", key=f"{base_key}_cancel_import", use_container_width=True):
                        st.session_state[import_key] = False
                        st.rerun()
                
                # Show import options if needed
                if st.session_state.get(f"{import_key}_show_options", False):
                    st.markdown("---")
                    st.markdown("**This vignette already has content. What would you like to do?**")
                    
                    col_opt1, col_opt2, col_opt3 = st.columns(3)
                    with col_opt1:
                        if st.button("📝 Replace Current", key=f"{base_key}_import_replace", use_container_width=True):
                            st.session_state[content_key] = st.session_state[f"{import_key}_pending"]
                            st.session_state[version_key] += 1
                            st.session_state[import_key] = False
                            st.session_state[f"{import_key}_pending"] = None
                            st.session_state[f"{import_key}_show_options"] = False
                            st.success("✅ File imported (replaced current content)!")
                            st.rerun()
                    
                    with col_opt2:
                        if st.button("➕ Append to Current", key=f"{base_key}_import_append", use_container_width=True):
                            current = st.session_state.get(content_key, "")
                            # Remove closing tags if any
                            current = current.replace('</p>', '')
                            new_content = current + st.session_state[f"{import_key}_pending"]
                            st.session_state[content_key] = new_content
                            st.session_state[version_key] += 1
                            st.session_state[import_key] = False
                            st.session_state[f"{import_key}_pending"] = None
                            st.session_state[f"{import_key}_show_options"] = False
                            st.success("✅ File imported (appended to current content)!")
                            st.rerun()
                    
                    with col_opt3:
                        if st.button("❌ Cancel Import", key=f"{base_key}_import_cancel_options", use_container_width=True):
                            st.session_state[f"{import_key}_pending"] = None
                            st.session_state[f"{import_key}_show_options"] = False
                            st.rerun()
        
        # Display spellcheck results if they exist (below the button row)
        if showing_results:
            result = st.session_state[spell_result_key]
            if "corrected" in result:
                st.markdown("---")
                st.markdown("### ✅ Suggested Corrections:")
                st.markdown(f'<div style="background-color: #f0f9ff; padding: 15px; border-radius: 8px; border-left: 4px solid #4CAF50;">{result["corrected"]}</div>', unsafe_allow_html=True)
                
                col_apply1, col_apply2, col_apply3 = st.columns([1, 1, 1])
                with col_apply2:
                    if st.button("📋 Apply Corrections", key=f"{base_key}_apply", type="primary", use_container_width=True):
                        corrected = result["corrected"]
                        if not corrected.startswith('<p>'):
                            corrected = f'<p>{corrected}</p>'
                        
                        st.session_state[content_key] = corrected
                        st.session_state[version_key] += 1  # Increment version to force remount
                        st.session_state[spell_result_key] = {"show": False}
                        st.success("✅ Corrections applied!")
                        st.rerun()
                    
                    if st.button("❌ Dismiss", key=f"{base_key}_dismiss", use_container_width=True):
                        st.session_state[spell_result_key] = {"show": False}
                        st.rerun()
            
            elif "message" in result:
                st.success(result["message"])
                if st.button("Dismiss", key=f"{base_key}_dismiss_msg"):
                    st.session_state[spell_result_key] = {"show": False}
                    st.rerun()
        
        # Display AI rewrite result if available
        if st.session_state.get(f"{base_key}_ai_result"):
            result = st.session_state[f"{base_key}_ai_result"]
            
            st.markdown("---")
            st.markdown(f"### {result.get('emoji', '✨')} AI Rewrite Result - {result['person']}")
            
            col_res1, col_res2 = st.columns(2)
            with col_res1:
                st.markdown("**📝 Original Version:**")
                with st.container():
                    st.markdown(f'<div style="background-color: #f0f0f0; padding: 15px; border-radius: 5px; border-left: 4px solid #ccc;">{result["original"]}</div>', unsafe_allow_html=True)
            
            with col_res2:
                st.markdown(f"**✨ Rewritten Version ({result['person']}):**")
                with st.container():
                    st.markdown(f'<div style="background-color: #e8f4fd; padding: 15px; border-radius: 5px; border-left: 4px solid #4a90e2;">{result["rewritten"]}</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.markdown("*This rewrite used your profile information to better capture your authentic voice.*")
            
            col_apply1, col_apply2, col_apply3 = st.columns(3)
            with col_apply1:
                if st.button("📋 Copy to Clipboard", key=f"{base_key}_ai_copy", use_container_width=True):
                    st.info("✅ Copied! Select the text above and press Ctrl+C")
            
            with col_apply2:
                if st.button("📝 Replace Original", key=f"{base_key}_ai_replace", type="primary", use_container_width=True):
                    # Wrap in paragraph tags if not present
                    new_content = result["rewritten"]
                    if not new_content.startswith('<p>'):
                        new_content = f'<p>{new_content}</p>'
                    
                    st.session_state[content_key] = new_content
                    st.session_state[version_key] += 1  # Increment version to force remount
                    del st.session_state[f"{base_key}_ai_result"]
                    st.session_state[f"{base_key}_show_ai_menu"] = False
                    st.rerun()
            
            with col_apply3:
                if st.button("🔄 Try Different Voice", key=f"{base_key}_ai_try_another", use_container_width=True):
                    del st.session_state[f"{base_key}_ai_result"]
                    st.session_state[f"{base_key}_show_ai_menu"] = True
                    st.rerun()
        
        # Preview section
        if st.session_state.get(f"{base_key}_show_preview", False) and st.session_state[content_key]:
            st.markdown("---")
            st.markdown("### 👁️ Preview")
            st.markdown(f"## {title or 'Untitled'}")
            st.markdown(f"**Theme:** {theme}  |  **Mood:** {mood}")
            st.markdown("---")
            st.markdown(st.session_state[content_key], unsafe_allow_html=True)
            
            if st.button("✕ Close Preview", key=f"{base_key}_close_preview"):
                st.session_state[f"{base_key}_show_preview"] = False
                st.rerun()
    
    def display_vignette_gallery(self, filter_by="all", on_select=None, on_edit=None, on_delete=None):
        if filter_by == "published":
            vs = [v for v in self.vignettes if not v.get("is_draft", True)]
        elif filter_by == "drafts":
            vs = [v for v in self.vignettes if v.get("is_draft", False)]
        else:
            vs = self.vignettes
        
        vs.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
        
        # Display success messages
        if st.session_state.get("publish_success"):
            st.success("🎉 Published successfully!")
            del st.session_state.publish_success
        if st.session_state.get("draft_success"):
            st.success("💾 Draft saved successfully!")
            del st.session_state.draft_success
        if st.session_state.get("edit_success"):
            st.success("✅ Changes saved successfully!")
            del st.session_state.edit_success
        if st.session_state.get("delete_success"):
            st.success("🗑️ Deleted successfully!")
            del st.session_state.delete_success
        
        if not vs:
            st.info("No vignettes yet. Click 'Create New Vignette' to start writing.")
            return
        
        for v in vs:
            with st.container():
                col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                
                with col1:
                    status_emoji = "📢" if not v.get("is_draft") else "📝"
                    status_text = "Published" if not v.get("is_draft") else "Draft"
                    st.markdown(f"### {status_emoji} {v['title']}  `{status_text}`")
                    st.markdown(f"*{v['theme']}*")
                    
                    content_preview = re.sub(r'<[^>]+>', '', v['content'])
                    if len(content_preview) > 100:
                        content_preview = content_preview[:100] + "..."
                    st.markdown(content_preview)
                    
                    date_str = datetime.fromisoformat(v.get('updated_at', v.get('created_at', ''))).strftime('%b %d, %Y')
                    st.caption(f"📝 {v['word_count']} words • Last updated: {date_str}")
                    if v.get('images'):
                        st.caption(f"📸 {len(v['images'])} image(s)")
                
                with col2:
                    if st.button("📖 Read", key=f"read_{v['id']}", use_container_width=True):
                        if on_select:
                            on_select(v['id'])
                
                with col3:
                    if st.button("✏️ Edit", key=f"edit_{v['id']}", use_container_width=True):
                        if on_edit:
                            on_edit(v['id'])
                
                with col4:
                    if st.button("🗑️ Delete", key=f"del_{v['id']}", use_container_width=True):
                        self.delete_vignette(v['id'])
                        st.session_state.delete_success = True
                        st.rerun()
                
                st.divider()
    
    def display_full_vignette(self, id, on_back=None, on_edit=None):
        v = self.get_vignette_by_id(id)
        if not v:
            return
        
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("← Back", use_container_width=True):
                if on_back:
                    on_back()
        
        status_emoji = "📢" if not v.get("is_draft") else "📝"
        status_text = "Published" if not v.get("is_draft") else "Draft"
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.caption(f"{status_emoji} **{status_text}**")
        with col2:
            st.caption(f"🎭 **{v.get('mood', 'Reflective')}**")
        with col3:
            st.caption(f"📝 **{v['word_count']} words**")
        with col4:
            created = datetime.fromisoformat(v.get('created_at', '')).strftime('%b %d, %Y')
            st.caption(f"📅 **Created: {created}**")
        
        st.markdown("---")
        st.markdown(f"# {v['title']}")
        st.markdown(f"*Theme: {v['theme']}*")
        st.markdown("---")
        st.markdown(v['content'], unsafe_allow_html=True)
        
        if v.get('images'):
            st.markdown("---")
            st.markdown("### 📸 Images")
            cols = st.columns(3)
            for i, img in enumerate(v['images']):
                with cols[i % 3]:
                    if img.get('base64'):
                        st.image(f"data:image/jpeg;base64,{img['base64']}", use_column_width=True)
                    elif img.get('path') and os.path.exists(img['path']):
                        st.image(img['path'], use_column_width=True)
                    if img.get('caption'):
                        st.caption(img['caption'])
        
        st.markdown("---")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("✏️ Edit", use_container_width=True, type="primary"):
                if on_edit:
                    on_edit(v['id'])
        
        with col2:
            if v.get("is_draft"):
                if st.button("📢 Publish Now", use_container_width=True):
                    v["is_draft"] = False
                    v["published_at"] = datetime.now().isoformat()
                    self.update_vignette(v["id"], v["title"], v["content"], v["theme"], v.get("mood"), v.get("images"))
                    st.success("🎉 Published!")
                    time.sleep(1)
                    st.rerun()
            else:
                if st.button("📝 Unpublish", use_container_width=True):
                    v["is_draft"] = True
                    self.update_vignette(v["id"], v["title"], v["content"], v["theme"], v.get("mood"), v.get("images"))
                    st.success("📝 Unpublished")
                    time.sleep(1)
                    st.rerun()
        
        with col3:
            if st.button("🗑️ Delete", use_container_width=True):
                self.delete_vignette(v['id'])
                st.session_state.delete_success = True
                if on_back:
                    on_back()
                st.rerun()
